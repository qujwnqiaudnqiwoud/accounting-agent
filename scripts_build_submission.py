from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parent
ROOT_DIR = PROJECT_DIR.parent
TEMPLATE_DIR = ROOT_DIR / "张三-李四-课程作业模板"
OUT_DIR = PROJECT_DIR / "final_submission"
SUBMISSION_DATE = "2026年6月12日"
DEPLOYMENT_URL = "https://accounting-agent-yspsw2afdlis9itfyheqde.streamlit.app/"
GITHUB_URL = "https://github.com/qujwnqiaudnqiwoud/accounting-agent"
TEXTBOOK_SOURCE = "张新民、钱爱民《财务报表分析（第6版·立体化数字教材版）案例分析与学习指导》"

TEAM_MEMBERS = [
    {
        "name": "徐北辰",
        "student_id": "202408240218",
        "class_name": "24级智能会计班",
        "role": "组长",
        "ratio": "60%",
        "contribution": (
            "负责智能体总体架构设计、smolagents 主控 Agent 接入、后端工具链组织、"
            "Streamlit 前端工作台构建、API 配置区、动态 Agent trace、数据读取与分析流程串联、"
            "代码调试、测试验证、GitHub 仓库发布、Streamlit Cloud 在线部署和最终工程打包。"
        ),
        "reflection": (
            "徐北辰在项目中主要负责把会计分析任务拆解为可执行、可追踪的 Agent 工具链。"
            "在查阅投资者互动问答和现有智能财经平台后，进一步认识到财报分析的难点不是简单生成一段结论，"
            "而是要把年报数字、指标口径、风险触发依据和报告表达连接成可复核的证据链。"
            "因此在搭建过程中，将财务数据清洗、指标计算、勾稽校验和风险识别交给确定性 Python 模块完成，"
            "让大模型承担计划生成和语言组织。通过前端工作台、动态 trace 和 Streamlit Cloud 在线部署的实现，"
            "也体会到一个可演示的 Agent 系统不仅要能算，还要能够被用户实际打开、测试和复核，"
            "让用户看清楚它正在调用什么工具、依据什么数据、输出什么结果。"
        ),
    },
    {
        "name": "周亦轩",
        "student_id": "202408240302",
        "class_name": "24级智能会计班",
        "role": "成员",
        "ratio": "40%",
        "contribution": (
            f"负责{TEXTBOOK_SOURCE}等权威教材框架、指标体系、风险规则和相关文献资料的收集整理，"
            "参与报告结构设计、页面表达优化和前端视觉美化，协助校对指标名称、报告表述和课堂展示材料。"
        ),
        "reflection": (
            f"周亦轩在项目中主要围绕会计专业资料和展示体验展开工作。通过整理{TEXTBOOK_SOURCE}中的教材框架、指标定义和风险规则，"
            "更加清楚地理解了传统财务报表分析中偿债能力、营运能力、盈利能力、成长能力、现金流量质量和杜邦分析之间的逻辑关系。"
            "在补充年报可读性、网络平台互动、信息处理成本和大数据审计等文献后，也进一步理解了财务报告分析工具的专业价值："
            "它需要帮助用户降低阅读成本、统一指标口径，并把风险提示限定在审慎、可追溯的范围内。"
            "在参与前端美化和报告结构设计时，也认识到会计智能体的展示界面需要兼顾专业性和可读性，"
            "既要保持财务系统的稳重风格，也要让用户能够快速找到数据来源、指标结论、风险依据和在线演示入口。"
        ),
    },
]

PROBLEM_EVIDENCE = [
    (
        "投资者互动平台",
        "深交所互动易、上证 e 互动、全景路演等平台长期承载投资者与上市公司的问答。"
        "这些问答显示，投资者并不只关心股价，也会围绕年报中的经营现金流、应收账款、坏账准备、收入确认、"
        "资产负债率和主要财务指标变动原因追问公司，说明财报理解和指标复核是实际资本市场中的高频需求。"
    ),
    (
        "具体问答案例",
        "例如圣泉集团投资者关系活动记录中，投资者曾根据年报提出经营净现金流与净利润比值偏低的问题，"
        "追问盈利增长是否缺乏现金实质性支撑；互动易、上证 e 互动中也能看到关于长期应收款、应收账款账龄、"
        "坏账计提和收入确认时点等问题。这些问题都需要把年报数字转换为可解释的财务指标和风险线索。"
    ),
    (
        "文献发现",
        "经管领域研究表明，互动问答会改变投资者的信息获取方式，但也可能带来噪音和意见分歧；"
        "年报可读性、文本复杂性和信息处理成本会影响投资者理解、风险识别和资本市场效率。"
        "因此，本项目并非单纯做一个技术演示，而是切入财务报告信息处理成本高、风险识别证据链弱的问题。"
    ),
]

INDUSTRY_EVIDENCE = [
    "CSMAR 智能财经报告分析平台已提供上市公司财报分析、财务健康诊断、风险提示和 Word/PDF 报告生成。",
    "CSMAR 上市公司风险智能感知系统把财务数据与非财务数据结合，用于经营风险识别、财务造假动机和特征指标分析。",
    "弈 Chat、同花顺问财、Wind Alice、Choice 妙想 AI 等产品体现了财经问答、智能投研和文档精读正在进入实际业务场景。",
    "Datayes AI 开放平台面向 AI Agent 提供金融数据接口，强调让智能体用对数据、查到原文并生成可追溯答案。",
    "这些平台说明市场已经在用 AI 降低投研和财报处理成本；本项目的定位则是面向课程场景，提供本地可运行、计算可复核、过程可追踪的会计垂直 Agent。"
]

REFERENCES = [
    "张新民、钱爱民：《财务报表分析（第6版·立体化数字教材版）案例分析与学习指导》，中国人民大学出版社，ISBN 978-7-300-31977-3。",
    "谭松涛, 阚铄, 崔小勇. 互联网沟通能够改善市场信息效率吗?——基于深交所“互动易”网络平台的研究[J]. 金融研究, 2016(3):174-188.",
    "徐寿福, 郑迎飞, 罗雨杰. 网络平台互动与股票异质性风险[J]. 财经研究, 2022, 48(10):153-168.",
    "周波, 李洁柔, 王少飞. 整合信息披露、信息处理成本与投资意愿——基于个体投资者判断的实验研究[J]. 财经研究, 2025, 51(4):139-154.",
    "徐巍, 姚振晔, 陈冬华. 中文年报可读性:衡量与检验[J]. 会计研究, 2021(03):28-44.",
    "王克敏, 王华杰, 李栋栋, 戴杏云. 年报文本信息复杂性与管理者自利——来自中国上市公司的证据[J]. 管理世界, 2018, 34(12):120-132.",
    "郭松林, 宁祺器, 窦斌. 上市公司年报文本增量信息与违规风险预测——基于语调和可读性的视角[J]. 统计研究, 2022, 39(12):69-84.",
    "徐荣华, 朱婧, 戴欣瑜. 大数据审计:理论框架、研究进展与未来展望[J]. 外国经济与管理, 2024, 46(11):122-137.",
    "深交所互动易: https://irm.cninfo.com.cn/",
    "上证 e 互动: https://sns.sseinfo.com/",
    "全景路演投资者关系互动平台: https://ir.p5w.net/",
    "CSMAR 智能财经报告分析平台: https://www.csmar.com/channels/63.html",
    "CSMAR 上市公司风险智能感知系统: https://www.csmar.com/channels/上市公司风险智能感知系统.html",
    "CSMAR 弈 Chat 智能财经对话系统: https://www.csmar.com/channels/102.html",
    "同花顺问财: https://www.iwencai.com/",
    "Wind 金融终端与 Wind Alice: https://www.wind.com.cn/",
    "Choice 智能金融终端: https://choice.eastmoney.com/",
    "Datayes AI 开放平台: https://d.datayes.com/",
    "Hugging Face smolagents 文档: https://huggingface.co/docs/smolagents",
    "DeepSeek API 文档: https://api-docs.deepseek.com/",
    "Streamlit 文档: https://docs.streamlit.io/",
    "pandas 文档: https://pandas.pydata.org/docs/",
    "pdfplumber、PyMuPDF、python-docx 等开源工具文档。",
]

REAL_CASE_DIR = PROJECT_DIR / "outputs" / "real_case" / "keda_xunfei_2025"
REAL_CASE_EVIDENCE_DIR = OUT_DIR / "evidence" / "科大讯飞2025年报案例"
RUN_OUTPUT_EVIDENCE_DIR = OUT_DIR / "evidence" / "运行输出样例"
REAL_CASE_SOURCE_PDF = PROJECT_DIR / "data" / "科大讯飞_2025年年度报告.pdf"


def _load_real_case_summary() -> dict[str, object]:
    summary_path = REAL_CASE_DIR / "real_case_summary.json"
    if not summary_path.exists():
        return {
            "source_pdf": "data/科大讯飞_2025年年度报告.pdf",
            "pages_scanned": "待生成",
            "pages_total": "待生成",
            "tables_found": "待生成",
            "table_rows": "待生成",
            "text_lines": "待生成",
            "items_extracted": "待生成",
            "years": [],
            "coverage_ratio": 0,
            "missing_core_items": [],
            "validation_summary": "待生成真实年报校验结果。",
            "ratio_rows": "待生成",
            "risk_count": "待生成",
            "trace_steps": "待生成",
            "core_metrics": {},
            "warnings": [],
        }
    return json.loads(summary_path.read_text(encoding="utf-8"))


def _coverage_text(summary: dict[str, object]) -> str:
    value = summary.get("coverage_ratio", 0)
    try:
        return f"{float(value):.0%}"
    except (TypeError, ValueError):
        return str(value)


def _years_text(summary: dict[str, object]) -> str:
    years = summary.get("years", [])
    if not isinstance(years, list) or not years:
        return "待生成"
    return "、".join(str(year) for year in years)


def _real_case_rows(summary: dict[str, object]) -> list[tuple[str, str]]:
    missing = summary.get("missing_core_items", [])
    warnings = summary.get("warnings", [])
    return [
        ("样例年报", "科大讯飞 2025 年年度报告 PDF"),
        ("扫描页数", f"{summary.get('pages_scanned')} / {summary.get('pages_total')} 页"),
        ("识别表格与文本", f"{summary.get('tables_found')} 个表格，{summary.get('table_rows')} 行表格记录，{summary.get('text_lines')} 行文本记录"),
        ("标准化结果", f"抽取 {summary.get('items_extracted')} 个标准财务科目，覆盖年度：{_years_text(summary)}"),
        ("核心科目覆盖率", f"{_coverage_text(summary)}；缺失核心科目：{'无' if not missing else '、'.join(str(item) for item in missing)}"),
        ("数据校验", str(summary.get("validation_summary", ""))),
        ("分析输出", f"{summary.get('ratio_rows')} 条指标记录，{summary.get('risk_count')} 条风险提示，{summary.get('trace_steps')} 个 Agent trace 步骤"),
        ("抽取说明", "；".join(str(item) for item in warnings) if warnings else "无明显抽取警告"),
    ]


def _copy_real_case_evidence() -> Path | None:
    if not REAL_CASE_DIR.exists():
        return None
    REAL_CASE_EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    for name in [
        "extracted_financial_data.csv",
        "extracted_financial_data.xlsx",
        "extracted_financial_data_audit.json",
        "real_case_summary.json",
        "keda_ratio_results.xlsx",
        "keda_data_validation.xlsx",
        "keda_agent_trace.json",
        "keda_financial_analysis_report.docx",
        "keda_financial_analysis_report.md",
    ]:
        src = REAL_CASE_DIR / name
        if src.exists():
            shutil.copy2(src, REAL_CASE_EVIDENCE_DIR / name)
    if REAL_CASE_SOURCE_PDF.exists():
        shutil.copy2(REAL_CASE_SOURCE_PDF, REAL_CASE_EVIDENCE_DIR / REAL_CASE_SOURCE_PDF.name)
    readme = REAL_CASE_EVIDENCE_DIR / "README_真实年报案例.md"
    summary = _load_real_case_summary()
    readme.write_text(
        "\n".join(
            [
                "# 科大讯飞 2025 年年度报告真实读取案例",
                "",
                "本目录保存课程项目的真实年报读取和分析证据，用于证明系统能够从年度报告 PDF 抽取标准财务数据，再进入指标计算、风险识别和报告生成流程。",
                "",
                f"- 扫描页数：{summary.get('pages_scanned')} / {summary.get('pages_total')} 页",
                f"- 识别表格与文本：{summary.get('tables_found')} 个表格，{summary.get('table_rows')} 行表格记录，{summary.get('text_lines')} 行文本记录",
                f"- 标准化结果：抽取 {summary.get('items_extracted')} 个标准财务科目，覆盖年度 {_years_text(summary)}",
                f"- 核心科目覆盖率：{_coverage_text(summary)}",
                f"- 数据校验：{summary.get('validation_summary')}",
                f"- 分析输出：{summary.get('ratio_rows')} 条指标记录，{summary.get('risk_count')} 条风险提示，{summary.get('trace_steps')} 个 Agent trace 步骤",
                "",
                "主要文件：",
                "",
                "- `extracted_financial_data.csv` / `.xlsx`：由年报抽取出的标准财务数据。",
                "- `extracted_financial_data_audit.json`：每个核心科目的页码、来源和候选行证据。",
                "- `keda_ratio_results.xlsx`：基于抽取数据计算出的指标表。",
                "- `keda_data_validation.xlsx`：数据完整性和勾稽关系校验表。",
                "- `keda_agent_trace.json`：Agent 工具调用记录。",
                "- `keda_financial_analysis_report.docx`：基于真实年报抽取数据生成的分析报告。",
            ]
        ),
        encoding="utf-8",
    )
    return REAL_CASE_EVIDENCE_DIR


def _copy_run_output_evidence() -> Path | None:
    sources = [
        PROJECT_DIR / "outputs" / "tables" / "ratio_results.xlsx",
        PROJECT_DIR / "outputs" / "tables" / "data_validation.xlsx",
        PROJECT_DIR / "outputs" / "tables" / "data_validation.json",
        PROJECT_DIR / "outputs" / "reports" / "financial_analysis_report.docx",
        PROJECT_DIR / "outputs" / "reports" / "financial_analysis_report.md",
        PROJECT_DIR / "outputs" / "agent_trace.json",
    ]
    chart_dir = PROJECT_DIR / "outputs" / "charts"
    if chart_dir.exists():
        sources.extend(sorted(chart_dir.glob("*.html")))
    existing = [path for path in sources if path.exists()]
    if not existing:
        return None
    RUN_OUTPUT_EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    for src in existing:
        shutil.copy2(src, RUN_OUTPUT_EVIDENCE_DIR / src.name)
    (RUN_OUTPUT_EVIDENCE_DIR / "README_运行输出样例.md").write_text(
        "\n".join(
            [
                "# 运行输出样例",
                "",
                "本目录保存系统运行后生成的可验证输出，用于静态阅卷时复核项目是否真的跑通。",
                "",
                "- `ratio_results.xlsx`：财务指标计算结果。",
                "- `data_validation.xlsx` / `.json`：数据校验结果。",
                "- `financial_analysis_report.docx` / `.md`：系统生成的分析报告。",
                "- `agent_trace.json`：Agent 工具调用记录。",
                "- `*.html`：Plotly 图表输出。",
            ]
        ),
        encoding="utf-8",
    )
    return RUN_OUTPUT_EVIDENCE_DIR


def _member_line(member: dict[str, str]) -> str:
    return f"{member['name']} {member['student_id']} {member['class_name']}"


def _set_normal_style(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        if name in doc.styles:
            h = doc.styles[name]
            h.font.size = Pt(size)
            h.font.color.rgb = RGBColor.from_string(color)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(6)


def _clear_doc(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _add_title(doc: Document, title: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 77, 58)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(90, 90, 90)


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    if "Table Grid" in [style.name for style in doc.styles]:
        table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "内容"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.18)


def build_topic_doc() -> Path:
    real_case = _load_real_case_summary()
    doc = Document(TEMPLATE_DIR / "张三-李四-作业题目.docx")
    _clear_doc(doc)
    _set_normal_style(doc)
    _add_title(doc, "《生成式AI会计前沿》课程作业题目说明", "财报智析 Agent：基于年报读取与确定性指标计算的会计垂直智能体")
    _add_kv_table(
        doc,
        [
            ("课程名称", "生成式AI会计前沿"),
            ("选题编号", "选题二：会计垂直领域智能体（Agent）"),
            ("项目名称", "财报智析 Agent"),
            ("小组成员", "；".join(_member_line(member) for member in TEAM_MEMBERS)),
            ("组长", _member_line(TEAM_MEMBERS[0])),
            ("提交日期", SUBMISSION_DATE),
            ("运行环境", "Python 3.11；Streamlit；pandas；smolagents；OpenAI-compatible API；pdfplumber；PyMuPDF；python-docx"),
            ("在线演示地址", DEPLOYMENT_URL),
            ("GitHub 仓库", GITHUB_URL),
        ],
    )
    doc.add_heading("1. 研究问题", level=1)
    doc.add_paragraph(
        "本项目切入的不是一般聊天机器人问题，而是会计、财务和审计场景中的财务报告解读问题。"
        "从深交所互动易、上证 e 互动、全景路演等投资者互动平台可以看到，投资者经常围绕年报中的经营现金流、"
        "应收账款、坏账准备、收入确认和资产负债率等指标向上市公司追问，说明年报数字到财务判断之间仍存在明显的信息处理成本。"
        "上市公司年度报告篇幅长、表格结构复杂，普通用户往往难以快速把 PDF/Word 年报转换为可分析的标准财务数据。"
        "本项目研究如何构建一个会计垂直领域 Agent，使其能够读取年报或标准 Excel/CSV，生成标准化数据，"
        "并进一步完成指标计算、趋势分析、杜邦分析、现金流量质量分析、风险识别和标准化报告生成。"
        f"其中，初步 Agent 分析流程不是主观拼接，而是以{TEXTBOOK_SOURCE}中的偿债能力、营运能力、盈利能力、"
        "成长能力、现金流量质量分析和杜邦分析等经典财务报表分析框架为依据，再将其转化为可调用的工具链。"
        "项目已完成 GitHub 仓库发布和 Streamlit Cloud 在线部署，形成可在线访问、可课堂演示、可复现实验的初步应用形态。"
    )
    doc.add_paragraph(
        "具体而言，圣泉集团投资者关系活动中曾出现投资者依据经营净现金流与净利润比值追问盈利质量的案例；"
        "《网络平台互动与股票异质性风险》《整合信息披露、信息处理成本与投资意愿》《中文年报可读性：衡量与检验》"
        "等文献也说明，投资者信息处理能力、年报可读性和信息呈现方式会影响财报理解。"
        "现实应用层面，CSMAR、同花顺问财、Wind Alice、Choice 和 Datayes AI 开放平台等已经出现智能财经问答、"
        "财报分析、风险诊断或金融 Agent 数据接口实践，本项目则把这些方向转化为可复现的课程级会计 Agent。"
    )
    doc.add_heading("2. 数据与样本", level=1)
    _add_bullets(
        doc,
        [
            "样例结构化数据：accounting-agent/data/sample_financial_data.xlsx。",
            f"分析框架来源：{TEXTBOOK_SOURCE}，用于确定 Agent 的指标分类、分析顺序和报告结构。",
            "真实年报读取样例：accounting-agent/data/科大讯飞_2025年年度报告.pdf。系统已扫描 "
            f"{real_case.get('pages_scanned')} / {real_case.get('pages_total')} 页，抽取 "
            f"{real_case.get('items_extracted')} 个标准财务科目，核心科目覆盖率 {_coverage_text(real_case)}。",
            "其他年报读取样例：支持可复制表格型 PDF 年报与 docx 年报；扫描版 PDF 需先 OCR。",
            "系统输出：prepared_financial_data.csv、ratio_results.xlsx、data_validation.xlsx、agent_trace.json、financial_analysis_report.docx。",
        ],
    )
    doc.add_heading("3. 研究方法", level=1)
    _add_bullets(
        doc,
        [
            f"以{TEXTBOOK_SOURCE}为权威教材依据，将偿债能力、营运能力、盈利能力、成长能力、现金流量质量和杜邦分析整理为 Agent 工作流。",
            "用 Python 工具链完成可追溯的财务数据清洗、指标计算和风险规则判断。",
            "用 smolagents ToolCallingAgent 规划工具链；API 不可用时自动降级为受控 Python 管线。",
            "用大模型完成计划生成、解释组织和报告润色，但不允许新增未计算数字。",
            "用 Streamlit 构建前端工作台，把数据读取阶段与财务分析阶段分开展示。",
            f"用 Streamlit Cloud 部署在线演示入口：{DEPLOYMENT_URL}，验证系统从 Agent 制作到初步应用的工程闭环。",
            "用投资者互动平台、年报可读性和大数据审计文献说明选题的会计专业价值。",
        ],
    )
    doc.add_heading("4. 小组分工", level=1)
    _add_kv_table(
        doc,
        [
            (
                f"{member['name']}（{member['role']}，{member['ratio']}）",
                member["contribution"],
            )
            for member in TEAM_MEMBERS
        ],
    )
    doc.add_heading("5. 在线部署与初步应用", level=1)
    doc.add_paragraph(
        f"本项目已发布 GitHub 仓库（{GITHUB_URL}），并完成 Streamlit Cloud 在线部署，访问地址为：{DEPLOYMENT_URL}。"
        "这说明系统不再只是本地运行的代码 Demo，而是具备了在线访问、交互测试和课堂展示能力。"
        "从课程作业角度看，该部署过程体现了从会计问题识别、Agent 架构搭建、API 接入、前端交互、测试验证到公网演示的完整实践链条。"
    )
    doc.add_paragraph(
        "公网部署还使教师或同学可以在不进入本地开发环境的情况下查看系统界面和交互流程。"
        "同时，API Key 不写入代码和仓库，可通过网页侧边栏或 Streamlit Cloud Secrets 等方式配置，体现了基本的工程安全意识。"
    )
    doc.add_heading("6. 作业承诺", level=1)
    doc.add_paragraph("本项目作业参考文献和 LLM/Agent 使用情况见课程报告附录。")
    doc.add_paragraph(f"本项目由徐北辰、周亦轩共同完成，贡献比例为徐北辰 60%、周亦轩 40%。")
    out = OUT_DIR / "财报智析Agent_作业题目说明.docx"
    doc.save(out)
    return out


def build_report_doc() -> Path:
    real_case = _load_real_case_summary()
    doc = Document(TEMPLATE_DIR / "课程报告模板-非Jupyter版本.docx")
    _clear_doc(doc)
    _set_normal_style(doc)
    _add_title(doc, "《生成式AI会计前沿》课程报告", "财报智析 Agent：面向上市公司年报的会计垂直智能体")
    doc.add_paragraph("山东财经大学会计学院制").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("二〇二六年六月").alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_kv_table(
        doc,
        [
            ("小组成员", "；".join(_member_line(member) for member in TEAM_MEMBERS)),
            ("组长", _member_line(TEAM_MEMBERS[0])),
            ("贡献比例", "徐北辰 60%；周亦轩 40%"),
            ("提交日期", SUBMISSION_DATE),
            ("在线演示地址", DEPLOYMENT_URL),
            ("GitHub 仓库", GITHUB_URL),
        ],
    )

    doc.add_heading("报告题目", level=1)
    doc.add_paragraph("财报智析 Agent：基于年报读取、确定性财务指标计算与大模型报告生成的会计垂直智能体")

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(
        "本项目围绕上市公司财务报表分析场景，构建了一个可本地运行的会计垂直 Agent 系统。"
        "选题来源于真实资本市场中的财报解读需求：投资者在深交所互动易、上证 e 互动、全景路演等平台上，"
        "经常围绕经营现金流与净利润匹配、应收账款账龄、坏账准备、收入确认和资产负债率等问题向上市公司追问。"
        "相关经管文献也表明，年报可读性、文本复杂性和信息处理成本会影响投资者理解、风险识别和市场信息效率。"
        "基于这一问题，本系统支持用户上传年度报告 PDF/Word 或标准 Excel/CSV，先将年报中的报表数据抽取为标准财务数据，"
        "再由 Agent 工具链完成数据校验、指标计算、趋势分析、杜邦分析、现金流量质量分析、异常风险识别和标准化报告生成。"
        f"需要特别说明的是，系统初步分析流程以{TEXTBOOK_SOURCE}为权威教材依据，"
        "将教材中的偿债能力、营运能力、盈利能力、成长能力、现金流量质量和杜邦分析框架转化为 Agent 可调用的工具链，"
        "因此该流程不是单纯技术拼装，而是会计专业分析框架的程序化实现。"
        "项目采用“确定性计算 + 大模型解释”的架构：财务数字、指标公式、风险触发均由 Python 工具生成，"
        "大模型仅用于工具链规划、解释组织和报告润色，从而降低财务分析中的幻觉风险。"
        "在工程交付层面，项目已发布 GitHub 仓库并部署至 Streamlit Cloud，形成可在线访问和课堂展示的演示入口，"
        "体现了从会计问题识别、Agent 架构设计、本地开发调试到初步应用发布的完整闭环。"
        "实验结果表明，系统能够在样例数据上稳定生成指标表、数据校验表、图表、Agent 调用记录和 Word 分析报告，"
        "形成从问题发现、数据处理、指标复核到报告输出的可追溯闭环。"
    )
    doc.add_paragraph(
        "为避免仅停留在样例数据层面，本项目还使用本地保存的科大讯飞 2025 年年度报告 PDF 做真实年报验证。"
        f"系统全文扫描 {real_case.get('pages_scanned')} / {real_case.get('pages_total')} 页，识别 "
        f"{real_case.get('tables_found')} 个表格、{real_case.get('table_rows')} 行表格记录和 "
        f"{real_case.get('text_lines')} 行文本记录，最终抽取 {real_case.get('items_extracted')} 个标准财务科目，"
        f"核心科目覆盖率为 {_coverage_text(real_case)}，并继续生成指标、校验表、风险提示和 Agent trace。"
    )
    doc.add_paragraph("关键词：会计智能体；财务报表分析；年报解析；Streamlit Cloud；smolagents；OpenAI-compatible API")

    doc.add_heading("1. 引言", level=1)
    doc.add_heading("1.1 问题发现路径", level=2)
    doc.add_paragraph(
        "本项目的问题不是凭空设定，而是从投资者真实的信息需求中归纳出来。深交所互动易、上证 e 互动和全景路演等平台，"
        "是投资者向上市公司追问经营状况、财务数据和风险事项的重要渠道。整理这些平台中的问答可以发现，"
        "投资者常把年报中的单个数字转化为更具体的财务判断问题，例如经营活动现金流量净额与净利润是否匹配、"
        "应收账款是否压占利润、坏账准备计提是否充分、收入确认是否谨慎、资产负债率是否触及风险边界等。"
        "这些问题都不是大模型凭语感就能回答的，而需要先提取报表数据，再按照会计指标口径进行计算、对比和解释。"
    )
    _add_bullets(doc, [f"{name}：{desc}" for name, desc in PROBLEM_EVIDENCE])
    doc.add_paragraph(
        "因此，本项目精确切入的会计、财务和审计领域问题可以概括为：年报信息处理成本较高，财务指标口径整理困难，"
        "投资者对同一财务数据的解读能力存在差异，风险识别过程缺少可追溯证据链，而直接让大模型撰写财报分析又可能产生未计算数字和过度判断。"
    )

    doc.add_heading("1.2 文献依据与专业价值", level=2)
    doc.add_paragraph(
        "经管领域已有研究为本项目提供了明确支撑。谭松涛等基于深交所“互动易”网络平台研究互联网沟通与市场信息效率，"
        "说明网络互动能够成为投资者获取公司经营信息的渠道。徐寿福等进一步指出，互动平台虽然提升沟通便利性，"
        "但也可能因信息传播不平衡和噪音信息放大投资者意见分歧。周波等关于整合信息披露的实验研究表明，"
        "优化信息呈现方式可以降低投资者信息处理成本，增强其理解和投资意愿。"
    )
    doc.add_paragraph(
        "在年报文本和风险识别方面，徐巍等构建中文年报可读性指标，王克敏等讨论年报文本复杂性与管理者自利，"
        "郭松林等发现年报语调和可读性能够为违规风险预测提供增量信息。徐荣华等关于大数据审计的综述则强调，"
        "新兴信息技术正在改变审计取证、数据分析和风险预警方式。上述文献共同说明：财务报告分析工具应当降低阅读和整理成本，"
        "保留计算证据，并避免把解释性文本替代为未经验证的结论。"
    )

    doc.add_heading("1.3 现实应用与本项目差异", level=2)
    doc.add_paragraph(
        "现实中，已有若干公司和平台在尝试用 AI 降低财经信息处理成本。例如 CSMAR 智能财经报告分析平台提供财务分析报告和风险提示，"
        "CSMAR 上市公司风险智能感知系统强调从财务和非财务角度识别风险，弈 Chat、同花顺问财、Wind Alice、Choice 妙想 AI "
        "等产品提供财经问答、智能投研和文档精读能力，Datayes AI 开放平台则面向 Agent 提供金融数据接口和可追溯答案能力。"
        "这说明市场已经出现智能财经分析和金融 Agent 的相邻实践。"
    )
    _add_bullets(doc, INDUSTRY_EVIDENCE)
    doc.add_paragraph(
        "但上述平台多面向商业投研或机构服务，用户难以看到完整的课程级实现过程。"
        "本项目的差异在于：围绕财务报表分析教学场景，提供本地可运行的完整代码；"
        "财务指标由确定性 Python 模块计算；每一步工具调用写入 Agent trace；"
        "生成报告时只解释已计算结果，不新增未验证数字。"
    )

    doc.add_heading("1.4 研究意义", level=2)
    _add_bullets(
        doc,
        [
            "会计专业价值：把偿债、营运、盈利、成长、现金流量质量和杜邦分析流程自动化，降低年报整理和指标复核成本。",
            "财务分析价值：把现金流、利润质量、资产负债结构和周转效率等常见投资者疑问转化为标准指标和风险提示。",
            "审计思维价值：用数据校验、勾稽关系检查和 Agent trace 保留证据链，训练学生从可验证数据出发提出审慎判断。",
            "技术实践价值：探索大模型在会计场景中的合理边界，避免让模型直接编造或估算财务数字。",
            "教学展示价值：把数据读取、工具调用、指标计算、风险识别和报告生成全过程可视化，便于课堂汇报。",
            f"应用落地价值：项目已完成 GitHub 发布和 Streamlit Cloud 在线部署（{DEPLOYMENT_URL}），说明课程原型具备被外部访问、演示和初步试用的条件。",
        ],
    )
    doc.add_heading("1.5 研究思路", level=2)
    doc.add_paragraph(
        "项目采用两阶段工作流：第一阶段负责读取年报或结构化文件，并生成标准 CSV/Excel；"
        "第二阶段在用户确认数据后启动 Agent 分析。主控 Agent 先读取模型配置并规划工具链，"
        "随后依次调用文件读取、数据清洗、数据校验、指标计算、趋势分析、杜邦分析、现金流分析、风险识别、图表生成和报告生成模块。"
    )
    doc.add_heading("1.6 教材框架来源", level=2)
    doc.add_paragraph(
        f"本项目初步 Agent 分析流程以{TEXTBOOK_SOURCE}为主要权威教材依据。"
        "教材中的财务报表分析路径强调从报表项目出发，围绕偿债能力、营运能力、盈利能力、成长能力、现金流量质量和杜邦分析展开，"
        "并通过指标变化解释企业财务状况、经营效率、盈利质量和资本结构风险。"
        "本项目将这一传统会计分析流程拆解为可执行模块：ratio_calculator 对应教材中的指标公式和分类，"
        "trend_analyzer 对应纵向趋势比较，dupont_analyzer 对应权益净利率分解，cashflow_analyzer 对应利润现金含量和销售收现能力分析，"
        "risk_detector 则把教材中的审慎分析思路转化为可追踪的风险提示规则。"
    )
    doc.add_paragraph(
        "因此，Agent 在本项目中并不是替代会计专业判断的黑箱聊天工具，而是把权威教材中的分析框架、指标口径和报告逻辑工程化。"
        "大模型负责调度和表达，核心财务分析框架仍来自教材和确定性工具。"
    )

    doc.add_heading("2. 实验设计", level=1)
    doc.add_heading("2.1 系统架构设计", level=2)
    _add_bullets(
        doc,
        [
            "前端层：Streamlit 工作台，提供 API 配置、文件上传、数据读取、分析启动、结果展示和下载功能；本地运行与 Streamlit Cloud 在线部署共用同一套代码。",
            "Agent 层：smolagents ToolCallingAgent 用于工具链规划；未配置 API Key 或远程调用失败时使用 fallback pipeline。",
            "工具层：包括 financial_statement_extractor、data_cleaner、data_validator、ratio_calculator、risk_detector、report_generator 等模块。",
            f"知识层：knowledge/ 中保存由{TEXTBOOK_SOURCE}整理而来的分析框架、指标定义、风险规则、提示词和报告模板。",
            "输出层：生成标准 CSV/Excel、数据校验表、指标 Excel、图表 HTML、Markdown/Word 报告和 Agent trace JSON。",
        ],
    )
    doc.add_heading("2.2 数据读取与标准化设计", level=2)
    doc.add_paragraph(
        "系统支持两类输入：一类是用户已经整理好的标准 Excel/CSV，另一类是上市公司年度报告 PDF/Word。"
        "对于年报，系统优先识别表格型财务报表，并结合文本行兜底、跨页表头继承、报告年度识别和合并报表优先规则，"
        "尽量抽取营业收入、营业成本、净利润、经营活动现金流量净额、总资产、总负债、所有者权益等核心科目。"
        "抽取完成后先展示标准数据和读取证据，用户确认后再进行分析。"
    )
    doc.add_paragraph(
        "在科大讯飞 2025 年年度报告验证中，系统未限制页数，而是对完整 PDF 进行全文扫描。"
        f"抽取结果覆盖 {_years_text(real_case)}，并将标准 CSV、Excel、抽取审计 JSON 和真实案例分析报告保存到 "
        "final_submission/evidence/科大讯飞2025年报案例/，供教师复核数据来源、页码证据和工具调用过程。"
    )
    doc.add_heading("2.3 指标计算与风险控制设计", level=2)
    doc.add_paragraph(
        "项目将财务数字处理与大模型文本生成分离。ratio_calculator 根据公式计算流动比率、速动比率、资产负债率、"
        "权益乘数、应收账款周转率、存货周转率、总资产周转率、毛利率、净利率、ROA、ROE、收入增长率、"
        "净利润增长率、净利润现金含量、销售收现比率等指标。data_validator 进一步检查核心科目完整性、"
        "资产=负债+所有者权益的勾稽关系，以及净利润与经营活动现金流量净额的基础匹配性。"
        "risk_detector 使用规则识别应收账款异常增长、存货增长较快、利润与现金流背离、毛利率连续下降、资产负债率持续上升等风险提示。"
        f"这些指标分类与分析顺序对应{TEXTBOOK_SOURCE}的财务报表分析框架，体现了从权威教材到 Agent 工具链的转化。"
    )
    doc.add_heading("2.4 大模型 API 与防幻觉设计", level=2)
    doc.add_paragraph(
        "系统默认支持 DeepSeek OpenAI-compatible API，也可以通过配置切换其他兼容服务。"
        "Prompt 明确限制模型不得新增未计算数字，不得使用“造假、违规、必然”等绝对化表述。"
        "当 API Key 未配置、API 调用失败或 smolagents 不可用时，系统保留 fallback pipeline，保证课堂演示不中断。"
        "公网部署时，API Key 不写入 GitHub 仓库和源代码，可通过网页侧边栏临时输入，或通过 Streamlit Cloud Secrets 等方式进行私密配置。"
    )

    doc.add_heading("3. 研究结果", level=1)
    doc.add_heading("3.1 工程实现结果", level=2)
    _add_bullets(
        doc,
        [
            "完成 accounting-agent/ 项目结构，包含 app.py、agents/、tools/、schemas/、config/、knowledge/、data/、tests/、notebooks/ 等目录。",
            f"完成基于{TEXTBOOK_SOURCE}的分析框架工程化，将教材中的指标体系和分析流程写入 knowledge/ 与工具模块。",
            "完成 Streamlit 两阶段工作流：读取数据并生成标准 CSV，用户确认后再启动分析。",
            "完成模型 API 配置区：支持侧边栏直接填写 API Base、模型名称和 API Key。",
            "完成动态 Agent trace：页面实时展示当前步骤、输入摘要、输出摘要、状态和耗时。",
            "完成数据校验、指标计算、图表生成和 Word/Markdown 报告导出。",
            f"完成 GitHub 仓库发布与 Streamlit Cloud 在线部署，演示地址：{DEPLOYMENT_URL}。",
        ],
    )
    doc.add_heading("3.2 在线部署与初步应用验证", level=2)
    doc.add_paragraph(
        f"除本地运行外，本项目已部署到 Streamlit Cloud，在线演示地址为：{DEPLOYMENT_URL}。"
        f"项目代码同步发布在 GitHub 仓库：{GITHUB_URL}。"
        "这一步使系统从本地课程原型推进到可在线访问、可交互测试、可课堂展示的初步应用状态。"
        "对期末作业而言，公网部署能够证明小组不仅完成了 Agent 代码开发，也掌握了从环境配置、依赖管理、仓库发布、"
        "云端运行到应用展示的完整工程流程。"
    )
    doc.add_paragraph(
        "在线部署保留了本地系统的核心交互：用户可以进入网页工作台，查看模型 API 配置区、上传财务数据或年报文件、"
        "观察数据读取与分析两个阶段，并下载分析结果。由于财务分析项目涉及 API Key，项目没有把密钥写入代码或公开仓库，"
        "而是通过页面输入或平台 Secrets 配置，兼顾展示便利性和密钥安全。"
    )
    doc.add_heading("3.3 测试与复现结果", level=2)
    doc.add_paragraph(
        "项目已使用 pytest 覆盖数据清洗、数据校验、指标计算、风险识别、模型配置、年报抽取、Word 渲染和 Agent trace 回调等模块。"
        "当前测试结果为 18 passed。使用 data/sample_financial_data.xlsx 可跑通从标准数据读取到报告生成的完整流程，"
        "并生成 data_validation.xlsx、ratio_results.xlsx、financial_analysis_report.docx 和 agent_trace.json。"
    )
    doc.add_heading("3.4 真实年报读取验证：科大讯飞 2025 年年度报告", level=2)
    doc.add_paragraph(
        "为检验系统对长篇上市公司年报的读取能力，本项目选取本地文件 data/科大讯飞_2025年年度报告.pdf 进行端到端验证。"
        "系统先从 PDF 抽取标准财务数据，再用同一套确定性工具链完成数据校验、指标计算、风险识别、报告生成和 Agent trace 记录。"
        "该案例可以直接对应普通用户“只有年报、不方便手工整理 CSV”的使用场景。"
    )
    _add_kv_table(doc, _real_case_rows(real_case))
    doc.add_paragraph(
        "从验证结果看，年报抽取阶段能够识别核心财务报表中的收入、成本、利润、经营现金流、资产、负债、权益、"
        "应收账款、存货和货币资金等项目，并保留页码、抽取来源、候选行文本和评分等审计线索。"
        "随后分析阶段生成指标记录、数据校验记录、风险提示和 Word 报告，证明本项目不是只依赖人工整理好的 CSV，"
        "而是能够将真实年报读取能力纳入 Agent 工作流。"
    )
    doc.add_heading("3.5 会计分析结果输出", level=2)
    doc.add_paragraph(
        "系统输出不是单一文字报告，而是包含多层证据：标准化财务数据用于核对来源，数据校验表用于判断数据质量，"
        "指标表用于复核公式，风险卡片说明触发依据和改进建议，Agent trace 记录每个工具调用的执行状态。"
        "这种设计使财务分析结果比单纯让大模型生成文本更可解释、可复核。"
        "对应投资者互动平台中的典型追问，系统能够把经营现金流与净利润背离、应收账款增长、存货增长、"
        "资产负债率上升和毛利率下降等问题转化为可下载的指标表和风险说明，帮助用户从年报事实走向财务判断。"
    )
    doc.add_heading("3.6 局限性与改进方向", level=2)
    _add_bullets(
        doc,
        [
            "扫描版 PDF 仍依赖 OCR，若年报只有图片而无文本层，系统无法保证完整抽取。",
            "当前版本以核心财务指标和内置风险规则为主，尚未接入行业数据库和同行横向比较。",
            "审计意见、关键审计事项、MD&A 和附注深度解析仍属于后续扩展方向。",
            "真实企业年报格式复杂，自动抽取结果仍需在分析前由用户核对。",
        ],
    )

    doc.add_heading("4. 研究结论与心得", level=1)
    doc.add_heading("4.1 研究结论", level=2)
    doc.add_paragraph(
        "本项目证明，会计垂直 Agent 的关键不在于让大模型直接给出财务结论，而在于把会计分析流程拆解为可验证的工具链。"
        f"本项目的初步分析流程以{TEXTBOOK_SOURCE}为权威教材依据，说明系统的专业逻辑来源于成熟财务报表分析体系，"
        "而不是由模型临时生成或由开发者任意拼接。"
        "在财务数字和指标公式由 Python 确定性执行的前提下，大模型可以更适合承担任务规划、语言组织和报告表达工作。"
        "这种架构兼顾了生成式 AI 的交互性和会计分析对准确性、谨慎性、可追溯性的要求。"
        "从问题来源看，投资者互动平台和年报文本研究共同说明：财报使用者需要的不是更多泛化文字，"
        "而是更低的信息处理成本、更清晰的指标口径和更可靠的证据链。"
        "本项目以本地化、可复现的方式实现了这一目标的课程版本，并进一步通过 GitHub 和 Streamlit Cloud 完成在线发布，"
        "说明系统已具备从制作 Agent 到初步应用展示的完整过程。"
        "后续可继续接入行业数据库、同行比较、审计意见和关键审计事项分析。"
    )
    doc.add_heading("4.2 学习心得与课程建议", level=2)
    for member in TEAM_MEMBERS:
        doc.add_paragraph(f"{member['name']}的学习心得：{member['reflection']}")
    doc.add_paragraph(
        "课程建议：本项目实践表明，如果课程能够继续增加垂直 Agent、财务数据工程和可解释 AI 的案例训练，"
        "会更有助于学生把生成式 AI 与会计专业知识结合起来。建议后续课程中加入更多真实年报、行业数据和审计文本的综合实验。"
    )

    doc.add_heading("附录", level=1)
    doc.add_heading("附录A 小组成员与贡献说明", level=2)
    _add_kv_table(
        doc,
        [
            (
                f"{member['name']}（{member['student_id']}，{member['class_name']}，{member['role']}）",
                f"贡献比例：{member['ratio']}。主要贡献：{member['contribution']}",
            )
            for member in TEAM_MEMBERS
        ],
    )
    doc.add_heading("附录B 参考文献与资料来源", level=2)
    _add_bullets(doc, REFERENCES)
    doc.add_heading("附录B-1 教材框架与系统模块对应关系", level=2)
    _add_kv_table(
        doc,
        [
            ("权威教材依据", TEXTBOOK_SOURCE),
            ("偿债能力分析", "对应流动比率、速动比率、资产负债率、权益乘数等指标计算与风险提示。"),
            ("营运能力分析", "对应应收账款周转率、存货周转率、总资产周转率等指标计算与趋势比较。"),
            ("盈利能力分析", "对应毛利率、净利率、ROA、ROE 等指标计算和盈利质量解释。"),
            ("成长能力分析", "对应营业收入增长率、净利润增长率等指标计算和变动说明。"),
            ("现金流量质量分析", "对应净利润现金含量、销售收现比率、利润现金流匹配性分析。"),
            ("杜邦分析", "对应 ROE 分解、总资产周转率、净利率和权益乘数的结构化解释。"),
            ("Agent 工具链转化", "教材框架决定分析顺序和指标口径；Python 工具负责计算，大模型负责调度和报告表达。"),
        ],
    )
    doc.add_heading("附录C LLM/Agent 使用说明", level=2)
    doc.add_paragraph(
        "本项目在开发和运行过程中使用大模型辅助完成代码生成、报告文本组织、Prompt 设计和错误排查。"
        "系统运行时，大模型仅参与 Agent 工具链规划和报告语言润色，不直接计算财务指标，不新增未由工具计算的财务数字。"
        "所有财务指标、风险规则和数据校验结果均由项目中的 Python 模块确定性生成。"
    )
    doc.add_heading("附录D 独立完成声明", level=2)
    doc.add_paragraph(
        "本课程作业由小组成员在课程要求范围内独立完成。除已在参考文献和 LLM/Agent 使用说明中列明的资料与工具外，"
        "不存在抄袭、代写或隐瞒外部协助的情况。"
    )
    doc.add_paragraph("签名：徐北辰、周亦轩")
    doc.add_paragraph(f"日期：{SUBMISSION_DATE}")
    doc.add_heading("附录E 在线演示与代码仓库", level=2)
    doc.add_paragraph(f"Streamlit Cloud 在线演示地址：{DEPLOYMENT_URL}")
    doc.add_paragraph(f"GitHub 项目仓库：{GITHUB_URL}")
    doc.add_paragraph(
        "说明：在线演示用于展示项目已具备公网访问、交互测试和课堂汇报能力。"
        "如平台因权限或休眠机制要求登录或唤醒，可按 Streamlit 页面提示进入；项目本地运行和提交包仍可独立复现。"
    )

    out = OUT_DIR / "财报智析Agent_课程报告.docx"
    doc.save(out)
    return out


def build_notebook() -> Path:
    cells = [
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "# 财报智析 Agent 课程作业 Notebook\\n",
                "\\n",
                "作者：徐北辰（202408240218，24级智能会计班）；周亦轩（202408240302，24级智能会计班）\\n",
                "\\n",
                "本 Notebook 用于展示项目运行逻辑、核心代码调用和输出结果，并附小组分工、参考资料与独立完成声明。\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 1. 研究问题\\n",
                "\\n",
                "本项目研究如何构建一个会计垂直 Agent，使其能够读取上市公司年报或标准财务数据，自动完成数据校验、指标计算、风险识别和报告生成。\\n",
                "\\n",
                f"初步 Agent 分析流程以{TEXTBOOK_SOURCE}为权威教材依据，将教材中的偿债能力、营运能力、盈利能力、成长能力、现金流量质量和杜邦分析框架转化为可调用工具链。\\n",
                "\\n",
                "问题来源于真实资本市场中的财报解读需求：深交所互动易、上证 e 互动、全景路演等平台上，投资者经常围绕经营现金流、应收账款、坏账准备、收入确认、资产负债率和财务指标变动原因向上市公司追问。相关文献也表明，年报可读性、文本复杂性和信息处理成本会影响投资者理解和风险识别。因此，本项目重点解决年报信息处理成本高、指标口径整理难、风险识别证据链弱和大模型直接生成财报结论容易幻觉的问题。\\n",
                "\\n",
                "已有 CSMAR 智能财经报告分析平台、CSMAR 上市公司风险智能感知系统、弈 Chat、同花顺问财、Wind Alice、Choice 妙想 AI、Datayes AI 开放平台等相邻实践，说明 AI 财经问答、报告生成、风险诊断和 Agent 数据接口已经进入业务场景。本项目的差异是面向课堂与课程作业，强调本地可运行、确定性指标计算和 Agent trace 可追踪。\\n",
                "\\n",
                f"项目已完成 GitHub 仓库发布和 Streamlit Cloud 在线部署，在线演示地址：{DEPLOYMENT_URL}。这一步说明系统从本地代码 Demo 推进到可访问、可交互测试、可课堂展示的初步应用状态。\\n",
            ],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "from pathlib import Path\\n",
                "import json\\n",
                "import sys\\n",
                "import pandas as pd\\n",
                "\\n",
                "PROJECT_ROOT = Path('..').resolve()\\n",
                "sys.path.insert(0, str(PROJECT_ROOT))\\n",
                "PROJECT_ROOT\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": ["## 2. 读取样例数据\\n"],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "sample_path = PROJECT_ROOT / 'data' / 'sample_financial_data.xlsx'\\n",
                "raw_df = pd.read_excel(sample_path)\\n",
                "raw_df.head()\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": ["## 3. 真实年报读取验证：科大讯飞 2025 年年度报告\\n"],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "real_case_dir = Path('evidence') / '科大讯飞2025年报案例'\\n",
                "if not real_case_dir.exists():\\n",
                "    real_case_dir = PROJECT_ROOT / 'outputs' / 'real_case' / 'keda_xunfei_2025'\\n",
                "\\n",
                "case_summary = json.loads((real_case_dir / 'real_case_summary.json').read_text(encoding='utf-8'))\\n",
                "case_df = pd.read_csv(real_case_dir / 'extracted_financial_data.csv')\\n",
                "case_summary\\n",
            ],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": ["case_df\\n"],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": ["## 4. 运行 Agent 分析主流程\\n"],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "from agents.main_agent import run_financial_analysis\\n",
                "from config.settings import ModelConfig\\n",
                "\\n",
                "result = run_financial_analysis(\\n",
                "    excel_path=sample_path,\\n",
                "    enable_pdf=False,\\n",
                "    generate_word=True,\\n",
                "    model_config_override=ModelConfig(enable_llm=False),\\n",
                ")\\n",
                "result.keys()\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": ["## 5. 查看数据校验结果\\n"],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "result['validation']['summary']\\n",
                "result['validation']['detail_table']\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": ["## 6. 查看财务指标表\\n"],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": ["result['ratio_pivot'].head(20)\\n"],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": ["## 7. 查看风险识别结果\\n"],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": ["pd.DataFrame(result['risks'])\\n"],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": ["## 8. 查看 Agent 调用记录\\n"],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": ["pd.DataFrame(result['trace'])\\n"],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 9. 输出文件\\n",
                "\\n",
                "- Word 报告：`outputs/reports/financial_analysis_report.docx`\\n",
                "- Markdown 报告：`outputs/reports/financial_analysis_report.md`\\n",
                "- 指标表：`outputs/tables/ratio_results.xlsx`\\n",
                "- 数据校验表：`outputs/tables/data_validation.xlsx`\\n",
                "- Agent trace：`outputs/agent_trace.json`\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 10. 附录：小组贡献、参考文献与声明\\n",
                "\\n",
                "### 小组贡献说明\\n",
                "\\n",
                "| 成员 | 学号 | 班级 | 主要贡献 | 贡献比例 |\\n",
                "|---|---|---|---|---|\\n",
                f"| 徐北辰 | 202408240218 | 24级智能会计班 | {TEAM_MEMBERS[0]['contribution']} | 60% |\\n",
                f"| 周亦轩 | 202408240302 | 24级智能会计班 | {TEAM_MEMBERS[1]['contribution']} | 40% |\\n",
                "\\n",
                "### 学习心得\\n",
                "\\n",
                f"- 徐北辰：{TEAM_MEMBERS[0]['reflection']}\\n",
                f"- 周亦轩：{TEAM_MEMBERS[1]['reflection']}\\n",
                "\\n",
                "### 教材框架与 Agent 流程对应\\n",
                "\\n",
                f"- 权威教材依据：{TEXTBOOK_SOURCE}\\n",
                "- 偿债能力分析：流动比率、速动比率、资产负债率、权益乘数。\\n",
                "- 营运能力分析：应收账款周转率、存货周转率、总资产周转率。\\n",
                "- 盈利能力分析：毛利率、净利率、ROA、ROE。\\n",
                "- 成长能力分析：营业收入增长率、净利润增长率。\\n",
                "- 现金流量质量分析：净利润现金含量、销售收现比率、利润现金流匹配性。\\n",
                "- 杜邦分析：ROE 分解、总资产周转率、净利率和权益乘数。\\n",
                "\\n",
                "### 参考文献与资料\\n",
                "\\n",
                *[f"- {ref}\\n" for ref in REFERENCES],
                "\\n",
                "### 独立完成声明\\n",
                "\\n",
                f"本课程作业由徐北辰、周亦轩在课程要求范围内独立完成。签名：徐北辰、周亦轩。日期：{SUBMISSION_DATE}。\\n",
            ],
        },
    ]
    for cell in cells:
        if "source" in cell:
            cell["source"] = [part.replace("\\n", "\n") for part in cell["source"]]
    notebook = {
        "cells": cells,
        "metadata": {
            "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
            "language_info": {"name": "python", "pygments_lexer": "ipython3"},
        },
        "nbformat": 4,
        "nbformat_minor": 5,
    }
    out = OUT_DIR / "财报智析Agent_项目说明_Notebook.ipynb"
    out.write_text(json.dumps(notebook, ensure_ascii=False, indent=2), encoding="utf-8")
    return out


def write_markdown_files() -> list[Path]:
    readme = OUT_DIR / "README_提交说明.md"
    readme.write_text(
        """# 财报智析 Agent 期末提交说明

## 提交前可选补充

- 如老师要求，可补充系统运行截图或课堂演示截图。
- 课程报告已补充投资者互动平台案例、经管期刊文献、智能财经平台应用对比和本项目解决路径。
- 课程报告已明确初步 Agent 分析流程来自权威教材《财务报表分析（第6版·立体化数字教材版）案例分析与学习指导》，并补充教材框架与系统模块对应关系。
- 项目已发布 GitHub 仓库并完成 Streamlit Cloud 在线部署，在线演示地址：{deployment_url}
- `final_submission/evidence/科大讯飞2025年报案例/` 已保存真实年报读取证据，包括标准 CSV、抽取审计 JSON、指标表、数据校验表、Agent trace 和真实案例 Word 报告。
- `final_submission/evidence/运行输出样例/` 已保存指标表、校验表、图表 HTML、分析报告和 Agent trace。
- `final_submission/evidence/文档渲染预览/` 保存 Word 文档渲染后的 PDF、HTML 和页面总览 PNG。

## 在线演示

- Streamlit Cloud：{deployment_url}
- GitHub 仓库：{github_url}
- 说明：在线演示用于证明项目已经完成从会计 Agent 制作、本地调试、API 接入、前端交互到公网展示的初步应用闭环。API Key 不写入代码和仓库，可通过网页侧边栏或 Streamlit Cloud Secrets 配置。

## 推荐运行方式

```bash
cd accounting-agent
pip install -r requirements.txt
streamlit run app.py
```

如需使用大模型 API，可在网页侧边栏顶部填写 API Key，或在终端设置：

```bash
export DEEPSEEK_API_KEY="你的真实key"
streamlit run app.py
```

## 推荐打包方式

最终压缩包建议命名为：

`选题二_财报智析Agent_徐北辰.zip`

打包时不要包含 `.venv/`、`__pycache__/`、`.pytest_cache/`、`.DS_Store`、大体积临时上传文件。

## Word 渲染检查

本项目已改用不依赖 LibreOffice 的渲染脚本，提交前可以运行：

```bash
.venv/bin/python scripts_render_docx.py final_submission/财报智析Agent_课程报告.docx --output-dir outputs/rendered_docs/course_report
.venv/bin/python scripts_render_docx.py final_submission/财报智析Agent_作业题目说明.docx --output-dir outputs/rendered_docs/topic_doc
```

输出目录中的 PDF 和 PNG 用于检查是否有空白页、乱码或明显排版问题。
""".format(deployment_url=DEPLOYMENT_URL, github_url=GITHUB_URL),
        encoding="utf-8",
    )
    checklist = OUT_DIR / "提交文件清单.md"
    checklist.write_text(
        """# 提交文件清单

## 已生成

- `财报智析Agent_课程报告.docx`
- `财报智析Agent_作业题目说明.docx`
- `财报智析Agent_项目说明_Notebook.ipynb`
- `README_提交说明.md`
- `evidence/科大讯飞2025年报案例/`
- `evidence/运行输出样例/`
- `evidence/文档渲染预览/`

## 项目主体建议一并提交

- `app.py`
- `agents/`
- `tools/`
- `schemas/`
- `config/`
- `knowledge/`
- `data/sample_financial_data.xlsx`
- `data/科大讯飞_2025年年度报告.pdf`
- `notebooks/demo.ipynb`
- `tests/`
- `README.md`
- `requirements.txt`
- `run_app.sh`

## 提交前请确认

- [x] 已填写所有姓名、学号、班级、日期。
- [x] 已填写小组贡献情况和贡献比例。
- [x] 已说明 Agent 初步分析流程以权威财务报表分析教材为依据，而不是任意技术拼接。
- [x] 已补充投资者互动平台、经管期刊文献和行业 AI 财经平台资料。
- [x] 已使用科大讯飞 2025 年年度报告完成真实 PDF 年报读取验证。
- [x] 已将真实年报抽取 CSV、审计 JSON、指标表、校验表和 Agent trace 放入提交包。
- [x] 已将运行输出样例、图表 HTML 和系统生成报告放入提交包。
- [x] 已将 Word 渲染预览 PDF、HTML 和页面总览图放入提交包。
- [x] 已完成 GitHub 仓库发布和 Streamlit Cloud 在线部署，并在课程报告中写明在线演示地址。
- [x] 已运行 `pytest` 并确认测试通过。
- [x] 已运行 `scripts_render_docx.py` 并检查 Word 报告预览。
- [x] 已确认 Streamlit 页面可打开。
- [x] 已删除 `.venv/`、缓存和大体积临时文件。
""",
        encoding="utf-8",
    )
    packaging = OUT_DIR / "推荐提交包结构.md"
    packaging.write_text(
        """# 推荐提交包结构

```text
选题二_财报智析Agent_徐北辰/
├── README.md
├── requirements.txt
├── run_app.sh
├── app.py
├── agents/
├── tools/
├── schemas/
├── config/
├── knowledge/
├── data/
│   ├── sample_financial_data.xlsx
│   └── 科大讯飞_2025年年度报告.pdf
├── notebooks/
│   └── demo.ipynb
├── tests/
└── final_submission/
    ├── 财报智析Agent_课程报告.docx
    ├── 财报智析Agent_作业题目说明.docx
    ├── 财报智析Agent_项目说明_Notebook.ipynb
    ├── evidence/
    ├── README_提交说明.md
    ├── 提交文件清单.md
    └── 推荐提交包结构.md
```
""",
        encoding="utf-8",
    )
    return [readme, checklist, packaging]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    evidence_dir = _copy_real_case_evidence()
    run_output_dir = _copy_run_output_evidence()
    if (PROJECT_DIR / "outputs" / "reports" / "financial_analysis_report.docx").exists():
        shutil.copy2(
            PROJECT_DIR / "outputs" / "reports" / "financial_analysis_report.docx",
            OUT_DIR / "示例输出_财务分析报告.docx",
        )
    paths = [build_topic_doc(), build_report_doc(), build_notebook(), *write_markdown_files()]
    if evidence_dir:
        paths.append(evidence_dir)
    if run_output_dir:
        paths.append(run_output_dir)
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()

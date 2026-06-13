from docx import Document

from scripts_render_docx import render_docx


def test_render_docx_without_libreoffice(tmp_path):
    docx_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("测试报告", level=1)
    doc.add_paragraph("这是一段用于验证 PyMuPDF 渲染链路的中文内容。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "结论"
    table.cell(1, 1).text = "渲染正常"
    doc.save(docx_path)

    result = render_docx(docx_path, tmp_path / "preview", dpi=72)

    assert result["page_count"] >= 1
    assert result["pdf_path"].exists()
    assert result["html_path"].exists()
    assert result["png_paths"][0].exists()

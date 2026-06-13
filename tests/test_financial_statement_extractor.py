from __future__ import annotations

import pandas as pd
from docx import Document

from tools.file_loader import load_financial_table
from tools.financial_statement_extractor import extract_financial_data_from_report


def test_load_financial_table_reads_csv(tmp_path):
    csv_path = tmp_path / "financial_data.csv"
    pd.DataFrame(
        {
            "项目": ["营业收入", "净利润"],
            "2023": [1200, 180],
            "2024": [1500, 210],
        }
    ).to_csv(csv_path, index=False, encoding="utf-8-sig")

    loaded = load_financial_table(csv_path)

    assert list(loaded.columns) == ["项目", "2023", "2024"]
    assert loaded.loc[0, "项目"] == "营业收入"


def test_extract_financial_data_from_docx_table(tmp_path):
    docx_path = tmp_path / "annual_report.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    header[0].text = "项目"
    header[1].text = "2022"
    header[2].text = "2023"
    header[3].text = "2024"

    rows = [
        ("营业收入", 1000, 1200, 1500),
        ("营业成本", 600, 720, 930),
        ("净利润", 120, 160, 210),
        ("经营活动现金流量净额", 130, 155, 230),
        ("资产总计", 3000, 3600, 4200),
        ("负债合计", 1200, 1500, 1750),
        ("所有者权益合计", 1800, 2100, 2450),
        ("流动资产合计", 900, 1100, 1400),
        ("流动负债合计", 500, 700, 850),
    ]
    for item, value_2022, value_2023, value_2024 in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = str(value_2022)
        cells[2].text = str(value_2023)
        cells[3].text = str(value_2024)
    doc.save(docx_path)

    result = extract_financial_data_from_report(docx_path, tmp_path)

    assert result["csv_path"].exists()
    assert result["xlsx_path"].exists()
    assert result["items_extracted"] >= 9
    assert result["years"] == [2022, 2023, 2024]
    extracted = result["dataframe"]
    revenue = extracted[extracted["项目"] == "营业收入"].iloc[0]
    assert revenue[2024] == 1500


def test_extract_financial_data_from_docx_text_lines(tmp_path):
    docx_path = tmp_path / "text_annual_report.docx"
    doc = Document()
    doc.add_paragraph("2024年年度报告")
    doc.add_paragraph("项目 2022 2023 2024")
    doc.add_paragraph("营业收入 1000 1200 1500")
    doc.add_paragraph("净利润 120 160 210")
    doc.save(docx_path)

    result = extract_financial_data_from_report(docx_path, tmp_path)

    extracted = result["dataframe"]
    revenue = extracted[extracted["项目"] == "营业收入"].iloc[0]
    assert revenue[2024] == 1500
    assert result["text_lines"] >= 3


def test_extract_financial_data_infers_current_and_previous_period_years(tmp_path):
    docx_path = tmp_path / "period_header_annual_report.docx"
    doc = Document()
    doc.add_paragraph("2024年年度报告")
    doc.add_paragraph("合并利润表")
    doc.add_paragraph("项目 本期金额 上期金额")
    doc.add_paragraph("营业收入 1500 1200")
    doc.add_paragraph("营业成本 930 720")
    doc.add_paragraph("净利润 210 160")
    doc.save(docx_path)

    result = extract_financial_data_from_report(docx_path, tmp_path)

    extracted = result["dataframe"]
    revenue = extracted[extracted["项目"] == "营业收入"].iloc[0]
    assert revenue[2024] == 1500
    assert revenue[2023] == 1200

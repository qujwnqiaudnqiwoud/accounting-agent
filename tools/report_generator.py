from __future__ import annotations

from pathlib import Path
from collections import Counter
from typing import Any

import pandas as pd

from config.settings import ModelConfig, PATHS
from tools.llm_client import polish_report


DISCLAIMER = "本结论基于用户提供的数据和教材分析框架生成，仅作为课程研究与财务分析参考。"


def _fmt(value: Any, percent: bool = False) -> str:
    if value is None or pd.isna(value):
        return "数据缺失"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    return f"{number:.2%}" if percent else f"{number:,.2f}"


def _latest_ratio(ratio_df: pd.DataFrame, ratio_name: str) -> tuple[int | None, float | None, bool]:
    part = ratio_df[(ratio_df["指标名称"] == ratio_name) & ratio_df["指标值"].notna()].sort_values("年份")
    if part.empty:
        return None, None, False
    row = part.iloc[-1]
    return int(row["年份"]), float(row["指标值"]), bool(row.get("是否百分比", False))


def _ratio_series(ratio_df: pd.DataFrame, ratio_name: str) -> list[tuple[int, float, bool]]:
    part = ratio_df[(ratio_df["指标名称"] == ratio_name) & ratio_df["指标值"].notna()].sort_values("年份")
    return [(int(row["年份"]), float(row["指标值"]), bool(row.get("是否百分比", False))) for _, row in part.iterrows()]


def _ratio_change_text(ratio_df: pd.DataFrame, ratio_name: str) -> str:
    series = _ratio_series(ratio_df, ratio_name)
    if not series:
        return f"{ratio_name}数据缺失。"
    latest_year, latest_value, is_percent = series[-1]
    latest_text = f"{ratio_name}{latest_year}年为{_fmt(latest_value, is_percent)}"
    if len(series) < 2:
        return latest_text + "。"
    prev_year, prev_value, _ = series[-2]
    if is_percent:
        change = (latest_value - prev_value) * 100
        direction = "上升" if change > 0 else "下降" if change < 0 else "持平"
        return f"{latest_text}，较{prev_year}年{direction}{abs(change):.2f}个百分点。"
    if prev_value == 0:
        return f"{latest_text}，因{prev_year}年基数为0，暂不计算同比变化。"
    change = latest_value / prev_value - 1
    direction = "上升" if change > 0 else "下降" if change < 0 else "持平"
    return f"{latest_text}，较{prev_year}年{direction}{abs(change):.2%}。"


def _item_series(clean_df: pd.DataFrame, item: str) -> list[tuple[int, float]]:
    if item not in clean_df.index:
        return []
    years = sorted([int(col) for col in clean_df.columns if isinstance(col, int) and pd.notna(clean_df.loc[item, col])])
    return [(year, float(clean_df.loc[item, year])) for year in years]


def _item_change_text(clean_df: pd.DataFrame, item: str, label: str) -> str:
    series = _item_series(clean_df, item)
    if not series:
        return f"{label}数据缺失。"
    latest_year, latest_value = series[-1]
    latest_text = f"{label}{latest_year}年为{_fmt(latest_value)}"
    if len(series) < 2:
        return latest_text + "。"
    prev_year, prev_value = series[-2]
    if prev_value == 0:
        return f"{latest_text}，因{prev_year}年基数为0，暂不计算同比变化。"
    change = latest_value / prev_value - 1
    direction = "增加" if change > 0 else "减少" if change < 0 else "持平"
    return f"{latest_text}，较{prev_year}年{direction}{abs(change):.2%}。"


def _trend_markdown(trend_df: pd.DataFrame, names: list[str], limit: int = 4) -> str:
    if trend_df.empty or "分析对象" not in trend_df.columns:
        return "暂无趋势判断。"
    rows = trend_df[trend_df["分析对象"].isin(names)].head(limit)
    if rows.empty:
        return "暂无趋势判断。"
    return "\n".join(
        f"- {row['分析对象']}：{row['趋势判断']}，{row['主要依据']}"
        for _, row in rows.iterrows()
    )


def _risk_profile(risks: list[dict[str, str]]) -> str:
    if not risks:
        return "内置规则未触发风险提示。"
    counts = Counter(risk.get("risk_level", "") for risk in risks)
    high = counts.get("高", 0)
    medium = counts.get("中", 0)
    low = counts.get("低", 0)
    leading = "、".join(risk["risk_name"] for risk in risks[:3])
    return f"本次共触发{len(risks)}条风险提示，其中高风险{high}条、中风险{medium}条、低风险{low}条；优先关注：{leading}。"


def _recommendation_markdown(risks: list[dict[str, str]], cashflow_result: dict[str, Any]) -> str:
    lines: list[str] = []
    if risks:
        for risk in risks[:4]:
            lines.append(f"- {risk['risk_name']}：{risk['suggestion']}")
    if cashflow_result.get("cashflow_quality") in {"风险较高", "需要关注"}:
        lines.append("- 现金流量质量：重点核查应收账款回款、存货占用和收入确认节奏。")
    if not lines:
        lines.append("- 未触发明显异常信号，建议继续补充行业均值和同行公司数据进行横向比较。")
    return "\n".join(lines)


def _ratio_markdown(ratio_df: pd.DataFrame) -> str:
    display_df = ratio_df.copy()
    display_df["指标值展示"] = display_df.apply(
        lambda row: "数据缺失"
        if pd.isna(row["指标值"])
        else f"{float(row['指标值']):.2%}"
        if bool(row["是否百分比"])
        else f"{float(row['指标值']):.2f}",
        axis=1,
    )
    pivot = display_df.pivot_table(index=["指标类别", "指标名称"], columns="年份", values="指标值展示", aggfunc="first")
    if pivot.empty:
        return "暂无可展示指标。"
    return pivot.reset_index().to_markdown(index=False)


def _risk_markdown(risks: list[dict[str, str]]) -> str:
    if not risks:
        return "未触发内置风险规则。"
    lines = []
    level_map = {"高": "高风险", "中": "中风险", "低": "低风险"}
    for risk in risks:
        level = level_map.get(risk["risk_level"], risk["risk_level"])
        lines.append(
            f"- **{risk['risk_name']}（{level}）**：{risk['evidence']} "
            f"{risk['explanation']} 建议：{risk['suggestion']}"
        )
    return "\n".join(lines)


def _validation_markdown(validation_result: dict[str, Any] | None) -> str:
    if not validation_result:
        return "未生成数据校验结果。"
    details = validation_result.get("details", [])
    if not details:
        return validation_result.get("summary", "未生成数据校验明细。")
    lines = [validation_result.get("summary", "数据校验完成。")]
    for row in details[:12]:
        lines.append(
            f"- {row.get('年份')}｜{row.get('校验项目')}｜{row.get('校验结果')}："
            f"{row.get('校验证据')} 建议：{row.get('处理建议')}"
        )
    if len(details) > 12:
        lines.append(f"- 其余 {len(details) - 12} 条校验明细见数据校验表。")
    return "\n".join(lines)


def build_report_draft(
    clean_df: pd.DataFrame,
    ratio_df: pd.DataFrame,
    trend_df: pd.DataFrame,
    dupont_result: dict[str, Any],
    cashflow_result: dict[str, Any],
    risks: list[dict[str, str]],
    validation_result: dict[str, Any] | None,
    trace: list[dict[str, Any]],
    pdf_info: dict[str, Any] | None,
    report_level: str,
) -> str:
    years = sorted([int(col) for col in clean_df.columns if isinstance(col, int)])
    latest_year = years[-1] if years else "最新年度"
    parsed_company_name = (pdf_info or {}).get("company_name")
    company_name = parsed_company_name if parsed_company_name and parsed_company_name != "未识别" else "用户上传公司"
    pdf_description = (
        f"PDF 年报解析识别公司名称为：{company_name}。"
        if pdf_info and pdf_info.get("status") == "success" and parsed_company_name and parsed_company_name != "未识别"
        else "PDF 已解析，但未识别公司名称，系统基于标准财务数据完成核心分析。"
        if pdf_info and pdf_info.get("status") == "success"
        else "未使用 PDF 年报或 PDF 解析未成功，系统基于标准财务数据完成核心分析。"
    )
    source_text = "标准财务数据"
    if pdf_info and pdf_info.get("status") == "success":
        source_text += "、PDF 年报文本"

    ratio_summary = []
    summary_names = [
        "营业收入增长率",
        "净利润增长率",
        "资产负债率",
        "净资产收益率（ROE）",
        "净利润现金含量",
    ]
    for name in summary_names:
        year, value, is_percent = _latest_ratio(ratio_df, name)
        ratio_summary.append(f"- {name}（{year or latest_year}）：{_fmt(value, is_percent)}")

    trace_lines = [
        f"- {item['order']}. {item['tool_name']}：{item['status']}，{item.get('output_summary', '')}"
        for item in trace
    ]
    scale_overview = "\n".join(
        [
            f"- {_item_change_text(clean_df, 'operating_revenue', '营业收入')}",
            f"- {_item_change_text(clean_df, 'net_profit', '净利润')}",
            f"- {_item_change_text(clean_df, 'net_operating_cash_flow', '经营活动现金流量净额')}",
            f"- {_item_change_text(clean_df, 'total_assets', '总资产')}",
        ]
    )
    balance_sheet_analysis = "\n".join(
        [
            _item_change_text(clean_df, "total_assets", "总资产"),
            _item_change_text(clean_df, "total_liabilities", "总负债"),
            _item_change_text(clean_df, "total_equity", "所有者权益"),
            _ratio_change_text(ratio_df, "资产负债率"),
            _ratio_change_text(ratio_df, "权益乘数"),
        ]
    )
    income_analysis = "\n".join(
        [
            _item_change_text(clean_df, "operating_revenue", "营业收入"),
            _item_change_text(clean_df, "operating_cost", "营业成本"),
            _item_change_text(clean_df, "net_profit", "净利润"),
            _ratio_change_text(ratio_df, "毛利率"),
            _ratio_change_text(ratio_df, "净利率"),
        ]
    )
    cashflow_analysis = "\n".join(
        [
            _item_change_text(clean_df, "net_operating_cash_flow", "经营活动现金流量净额"),
            _ratio_change_text(ratio_df, "净利润现金含量"),
            _ratio_change_text(ratio_df, "销售收现比率"),
            _ratio_change_text(ratio_df, "经营现金流量收入比率"),
            str(cashflow_result.get("cashflow_summary", "")),
        ]
    )
    solvency_analysis = "\n".join(
        [
            _ratio_change_text(ratio_df, "流动比率"),
            _ratio_change_text(ratio_df, "速动比率"),
            _ratio_change_text(ratio_df, "资产负债率"),
            _ratio_change_text(ratio_df, "权益乘数"),
        ]
    )
    operating_analysis = "\n".join(
        [
            _ratio_change_text(ratio_df, "应收账款周转率"),
            _ratio_change_text(ratio_df, "存货周转率"),
            _ratio_change_text(ratio_df, "总资产周转率"),
            _trend_markdown(trend_df, ["应收账款", "存货", "总资产"], limit=3),
        ]
    )
    profitability_analysis = "\n".join(
        [
            _ratio_change_text(ratio_df, "毛利率"),
            _ratio_change_text(ratio_df, "净利率"),
            _ratio_change_text(ratio_df, "总资产报酬率（ROA）"),
            _ratio_change_text(ratio_df, "净资产收益率（ROE）"),
            dupont_result.get("dupont_summary", "杜邦分析数据不足。"),
        ]
    )
    growth_analysis = "\n".join(
        [
            _ratio_change_text(ratio_df, "营业收入增长率"),
            _ratio_change_text(ratio_df, "净利润增长率"),
            _ratio_change_text(ratio_df, "总资产增长率"),
            _ratio_change_text(ratio_df, "所有者权益增长率"),
            "若收入增长与净利润、经营活动现金流量净额变化不同步，应优先检查毛利率、费用率、营运资金占用和非经常性损益影响。",
        ]
    )

    report = f"""# {company_name} 财务报表分析报告

**报告期间**：{years[0] if years else "数据缺失"} 年度至 {latest_year} 年度  
**报告详细程度**：{report_level}  
**分析工具**：财报智析 Agent  
**数据来源**：{source_text}

> {DISCLAIMER}

## 一、公司基本情况

本报告基于已确认的标准财务数据生成。{pdf_description}

## 二、分析数据来源与方法说明

系统按照教材框架执行数据读取、项目标准化、指标计算、趋势分析、杜邦分析、现金流量质量分析、风险识别和报告生成。财务数字由 Python 工具根据标准财务数据计算，大模型只用于调度说明和文字组织；若 PDF 已解析，年报文本仅作为补充背景，不替代结构化财务数据。

## 三、数据质量与勾稽校验

{_validation_markdown(validation_result)}

## 四、主要财务数据概览

{chr(10).join(ratio_summary)}

{scale_overview}

核心趋势：

{_trend_markdown(trend_df, ["营业收入", "净利润", "经营活动现金流量净额", "总资产"], limit=4)}

## 五、资产负债表分析

{balance_sheet_analysis}

资产负债表分析重点不只看规模扩张，还要看负债扩张是否快于权益积累。若资产负债率与权益乘数同步上升，说明企业更多依赖负债或杠杆支撑资产规模，需要进一步拆分金融性负债、经营性负债和债务期限结构。

## 六、利润表分析

{income_analysis}

利润表分析重点看收入增长是否转化为毛利和净利润。若营业收入增长而毛利率或净利率下降，说明增长质量可能受到成本、价格、费用或产品结构影响；若净利润增长弱于收入增长，应进一步核查期间费用、资产减值损失和非经常性损益。

## 七、现金流量表分析

{cashflow_analysis}

现金流量表分析重点看利润是否有经营现金流支撑。净利润现金含量低于100%并持续下降时，通常意味着利润确认与现金回收存在时间差或营运资金占用加重，需要结合应收账款、存货和合同负债变化判断。

## 八、偿债能力分析

{solvency_analysis}

偿债能力分析同时覆盖短期流动性和长期资本结构。流动比率、速动比率下降时，应关注流动负债到期压力和流动资产质量；资产负债率、权益乘数上升时，应进一步判断杠杆扩张是否带来足够的盈利和现金流回报。

## 九、营运能力分析

{operating_analysis}

营运能力分析重点看资金占用效率。应收账款周转率下降通常意味着回款周期拉长；存货周转率下降提示存货消化速度变慢；总资产周转率下降则说明资产投入转化为收入的效率减弱。

## 十、盈利能力分析

{profitability_analysis}

盈利能力分析需要同时观察利润率、资产报酬和股东权益回报。净资产收益率（ROE）不能单独解释盈利改善，必须结合净利率、总资产周转率和权益乘数判断：若ROE变化主要由权益乘数驱动，应关注杠杆风险；若由净利率驱动，则应继续拆分毛利、费用和减值因素。

## 十一、成长能力分析

{growth_analysis}

成长能力分析关注“规模、利润、现金流”三者是否匹配。收入增长不能单独证明经营质量改善，只有净利润和经营活动现金流量净额同步改善，成长才更具可持续性。

## 十二、杜邦综合分析

{dupont_result.get("dupont_summary", "杜邦分析数据不足。")}

## 十三、现金流量质量与盈余质量分析

现金流量质量评价：**{cashflow_result.get("cashflow_quality", "数据不足")}**。  
{cashflow_result.get("cashflow_summary", "")}

## 十四、异常信号与风险提示

{_risk_profile(risks)}

{_risk_markdown(risks)}

## 十五、综合评价与改进建议

综合判断应优先围绕盈利质量、现金流量匹配、偿债压力和营运资金占用展开。建议按以下顺序核查：

{_recommendation_markdown(risks, cashflow_result)}

## 十六、附录：指标公式、数据来源、Agent 工具调用记录

### 主要指标表

{_ratio_markdown(ratio_df)}

### Agent 工具调用记录

{chr(10).join(trace_lines)}
"""
    return report


def export_docx(markdown_text: str, output_path: Path) -> Path | None:
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped.replace("**", ""))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_report(
    clean_df: pd.DataFrame,
    ratio_df: pd.DataFrame,
    trend_df: pd.DataFrame,
    dupont_result: dict[str, Any],
    cashflow_result: dict[str, Any],
    risks: list[dict[str, str]],
    validation_result: dict[str, Any] | None,
    trace: list[dict[str, Any]],
    pdf_info: dict[str, Any] | None,
    model_config: ModelConfig,
    report_level: str = "标准版",
    generate_word: bool = True,
) -> dict[str, Any]:
    draft = build_report_draft(
        clean_df,
        ratio_df,
        trend_df,
        dupont_result,
        cashflow_result,
        risks,
        validation_result,
        trace,
        pdf_info,
        report_level,
    )
    validation_summary = (validation_result or {}).get("summary", "未生成数据校验结果")
    source_summary = (
        f"指标数量：{len(ratio_df)}；风险数量：{len(risks)}；"
        f"现金流评价：{cashflow_result.get('cashflow_quality')}；{validation_summary}"
    )
    polished = polish_report(model_config, draft, source_summary) if model_config.is_configured else None
    markdown_text = polished or draft
    if DISCLAIMER not in markdown_text:
        markdown_text += f"\n\n> {DISCLAIMER}\n"

    PATHS.reports_dir.mkdir(parents=True, exist_ok=True)
    md_path = PATHS.reports_dir / "financial_analysis_report.md"
    md_path.write_text(markdown_text, encoding="utf-8")
    docx_path = None
    if generate_word:
        docx_path = export_docx(markdown_text, PATHS.reports_dir / "financial_analysis_report.docx")
    return {
        "markdown": markdown_text,
        "markdown_path": md_path,
        "word_path": docx_path,
        "used_llm": bool(polished),
    }

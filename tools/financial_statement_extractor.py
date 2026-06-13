from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd

from schemas.financial_schema import REQUIRED_CORE_ITEMS, STANDARD_ITEMS
from tools.data_cleaner import clean_financial_data


DISPLAY_ITEMS = {
    "operating_revenue": "营业收入",
    "operating_cost": "营业成本",
    "net_profit": "净利润",
    "net_operating_cash_flow": "经营活动现金流量净额",
    "cash_received_from_sales": "销售商品、提供劳务收到的现金",
    "total_assets": "总资产",
    "total_liabilities": "总负债",
    "total_equity": "所有者权益",
    "current_assets": "流动资产合计",
    "current_liabilities": "流动负债合计",
    "accounts_receivable": "应收账款",
    "inventory": "存货",
    "cash": "货币资金",
    "asset_impairment_loss": "资产减值损失",
    "non_recurring_pnl": "非经常性损益",
}

EXTRA_ALIASES = {
    "operating_revenue": ["一、营业总收入", "其中：营业收入", "营业收入合计"],
    "operating_cost": ["营业成本合计", "其中：营业成本"],
    "net_profit": ["五、净利润", "归属于母公司所有者的净利润", "归属于上市公司股东的净利润"],
    "net_operating_cash_flow": [
        "经营活动产生的现金流量净额",
        "经营活动现金流量净额小计",
        "经营活动产生的现金流量净额小计",
    ],
    "cash_received_from_sales": ["销售商品和提供劳务收到的现金", "销售商品、提供劳务收到的现金"],
    "total_assets": ["资产总额", "资产总计"],
    "total_liabilities": ["负债总额", "负债总计"],
    "total_equity": [
        "股东权益总计",
        "股东权益合计",
        "归属于母公司股东权益合计",
        "归属于母公司所有者权益合计",
        "所有者权益(或股东权益)合计",
        "所有者权益（或股东权益）合计",
    ],
    "current_assets": ["流动资产总计", "流动资产合计", "动资产合计"],
    "current_liabilities": ["流动负债总计", "流动负债合计", "动负债合计"],
    "accounts_receivable": ["应收账款账面价值", "应收账款净额"],
    "inventory": ["存货账面价值", "存货净额"],
    "cash": ["货币资金", "现金及现金等价物余额", "现金及现金等价物"],
}

EXCLUDED_ROW_PATTERNS = [
    "负债和所有者权益总计",
    "负债及所有者权益总计",
    "负债和股东权益总计",
    "负债及股东权益总计",
]

BALANCE_SHEET_ITEMS = {
    "total_assets",
    "total_liabilities",
    "total_equity",
    "current_assets",
    "current_liabilities",
    "accounts_receivable",
    "inventory",
    "cash",
}
INCOME_STATEMENT_ITEMS = {"operating_revenue", "operating_cost", "net_profit", "asset_impairment_loss"}
CASHFLOW_ITEMS = {"net_operating_cash_flow", "cash_received_from_sales"}

PDF_TABLE_SETTINGS: list[dict[str, Any] | None] = [
    None,
    {
        "vertical_strategy": "lines",
        "horizontal_strategy": "lines",
        "snap_tolerance": 3,
        "join_tolerance": 3,
        "intersection_tolerance": 5,
    },
    {
        "vertical_strategy": "text",
        "horizontal_strategy": "text",
        "text_x_tolerance": 2,
        "text_y_tolerance": 3,
        "snap_tolerance": 3,
        "join_tolerance": 3,
        "min_words_vertical": 2,
        "min_words_horizontal": 1,
    },
]

STATEMENT_TITLES = [
    "合并资产负债表",
    "母公司资产负债表",
    "合并利润表",
    "母公司利润表",
    "合并现金流量表",
    "母公司现金流量表",
]


@dataclass
class SourceRow:
    cells: list[str]
    page: int | None
    source: str
    context: str = ""

    @property
    def row_text(self) -> str:
        return "".join(self.cells)


@dataclass
class CandidateValue:
    value: float
    score: int
    page: int | None
    source: str
    row_text: str


def _max_pdf_pages() -> int | None:
    raw = os.getenv("FIN_AGENT_REPORT_MAX_PAGES", "0").strip()
    try:
        pages = int(raw)
    except ValueError:
        return None
    return pages if pages > 0 else None


def _normalize(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\n", "").replace(" ", "").replace("\u3000", "")
    text = text.replace("（", "(").replace("）", ")")
    text = text.replace("－", "-").replace("—", "-").replace("−", "-")
    return text.strip()


def _display_text(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\u3000", " ").replace("（", "(").replace("）", ")")
    text = text.replace("－", "-").replace("—", "-").replace("−", "-")
    return re.sub(r"\s+", " ", text).strip()


def _number(value: Any) -> float | None:
    if value is None:
        return None
    text = _display_text(value)
    if not text or text in {"-", "--"}:
        return None
    if "%" in text or "％" in text:
        return None
    text = text.replace(",", "").replace("，", "")
    text = re.sub(r"(人民币|单位|万元|千元|百万元|亿元|元|：|:)", "", text)
    negative = text.startswith("(") and text.endswith(")")
    text = text.strip("()")
    amount_strings = _amount_strings(text)
    if not amount_strings:
        return None
    try:
        number = float(amount_strings[0].replace(",", "").replace("，", "").strip("()"))
    except ValueError:
        return None
    return -number if negative else number


def _separate_joined_amounts(text: str) -> str:
    separated = text
    separated = re.sub(r"(\.\d{2})(?=-?\d{1,3}[,，]\d{3})", r"\1 ", separated)
    separated = re.sub(r"(\.\d{2})(?=-?\d+\.\d{2})", r"\1 ", separated)
    separated = re.sub(r"(%|％)(?=-?\d)", r"\1 ", separated)
    return separated


def _amount_strings(value: Any) -> list[str]:
    text = _display_text(value)
    if not text:
        return []
    text = _separate_joined_amounts(text)
    pattern = r"\(?-?(?:\d{1,3}(?:[,，]\d{3})+|\d+)(?:\.\d{1,2})?\)?(?:%|％)?"
    amounts: list[str] = []
    for match in re.finditer(pattern, text):
        token = match.group(0)
        if "%" in token or "％" in token:
            continue
        next_char = text[match.end() : match.end() + 1]
        prev_char = text[match.start() - 1 : match.start()] if match.start() > 0 else ""
        if next_char in {"、", ".", "．", ")"} and (not prev_char or prev_char.isspace()):
            continue
        amounts.append(token)
    return amounts


def _numbers_from_cells(cells: list[Any]) -> list[float]:
    numbers: list[float] = []
    for cell in cells:
        for token in _amount_strings(cell):
            value = _number(token)
            if value is not None:
                numbers.append(value)
    return numbers


def _alias_map() -> dict[str, str]:
    aliases: dict[str, str] = {}
    for canonical, names in STANDARD_ITEMS.items():
        aliases[_normalize(canonical)] = canonical
        for name in names:
            aliases[_normalize(name)] = canonical
    for canonical, names in EXTRA_ALIASES.items():
        for name in names:
            aliases[_normalize(name)] = canonical
    return aliases


def _match_item(row_text: str, aliases: dict[str, str]) -> str | None:
    normalized = _normalize(row_text)
    if any(pattern in normalized for pattern in EXCLUDED_ROW_PATTERNS):
        return None
    if any(
        pattern in normalized
        for pattern in [
            "现金及现金等价物净增加额",
            "现金及现金等价物净变动",
            "现金的期末余额",
            "现金的期初余额",
            "应收账款坏账",
            "应收账款组合",
            "存货的减少",
            "存货跌价",
        ]
    ):
        return None
    if "营业总成本" in normalized:
        return None
    if "其他流动资产" in normalized or "其他流动负债" in normalized:
        return None
    if "非流动资产" in normalized or "非流动负债" in normalized:
        return None
    matches = [(alias, canonical) for alias, canonical in aliases.items() if alias and alias in normalized]
    if not matches:
        return None
    matches.sort(key=lambda item: len(item[0]), reverse=True)
    canonical = matches[0][1]
    if canonical == "total_equity" and not any(token in normalized for token in ["合计", "总计"]):
        return None
    if canonical == "total_liabilities" and not any(token in normalized for token in ["合计", "总计", "总额"]):
        return None
    if canonical == "total_assets" and not any(token in normalized for token in ["总计", "总额"]):
        return None
    if canonical in {"current_assets", "current_liabilities"} and not any(token in normalized for token in ["合计", "总计"]):
        return None
    return canonical


def _infer_report_year_from_text(text: str) -> int | None:
    patterns = [
        r"(20\d{2})\s*年\s*年度报告",
        r"(20\d{2})\s*年度报告",
        r"(20\d{2})\s*年\s*年度财务报告",
        r"(20\d{2})\s*年\s*1\s*月\s*1\s*日.*?(20\d{2})\s*年\s*12\s*月\s*31\s*日",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "")
        if match:
            return int(match.group(1))
    return None


def _infer_report_year(path: Path, rows: list[SourceRow]) -> int | None:
    from_name = _infer_report_year_from_text(path.name)
    if from_name:
        return from_name
    for row in rows[:300]:
        year = _infer_report_year_from_text(row.context or row.row_text)
        if year:
            return year
    return None


def _context_years(context: str, report_year: int | None = None) -> list[int]:
    years = {int(match.group(1)) for match in re.finditer(r"(20\d{2})", context or "")}
    if report_year:
        years = {year for year in years if report_year - 6 <= year <= report_year}
    return sorted(years)


def _years_in_row(row: list[Any], context: str = "", report_year: int | None = None) -> dict[int, int]:
    years: dict[int, int] = {}
    for idx, cell in enumerate(row):
        for match in re.finditer(r"(20\d{2})", str(cell or "")):
            year = int(match.group(1))
            if report_year and year > report_year:
                continue
            if report_year and year < report_year - 6:
                continue
            years[idx] = year
            break

    if years:
        return years

    context_report_year = report_year or _infer_report_year_from_text(context)
    if not context_report_year:
        context_years = _context_years(context, report_year)
        context_report_year = max(context_years) if context_years else None
    if not context_report_year:
        return years

    current_year = context_report_year
    previous_year = context_report_year - 1
    for idx, cell in enumerate(row):
        text = _normalize(cell)
        if any(token in text for token in ["本期", "本年", "期末", "本报告期", "本期金额", "期末余额"]):
            years[idx] = current_year
        elif any(token in text for token in ["上期", "上年", "期初", "上年同期", "上期金额", "期初余额"]):
            years[idx] = previous_year
    return years


def _split_text_line(line: str) -> list[str]:
    text = _display_text(line)
    if not text:
        return []
    if re.search(r"20\d{2}", text) or any(token in text for token in ["本期", "上期", "本年", "上年", "期末", "期初"]):
        parts = [part.strip() for part in re.split(r"\s+", text) if part.strip()]
        return parts
    parts = [part.strip() for part in re.split(r"\s{2,}|\t+", text) if part.strip()]
    if len(parts) > 1:
        return parts
    numbers = _amount_strings(text)
    if len(numbers) >= 1:
        first_number = text.find(numbers[0])
        item = text[:first_number].strip()
        if item:
            return [item, *numbers]
    return [text]


def _context_score(context: str) -> int:
    normalized = _normalize(context)
    score = 0
    is_note_context = any(token in normalized for token in ["项目注释", "财务报表附注", "会计政策", "会计估计", "补充资料"])
    if any(token in normalized for token in ["合并资产负债表", "合并利润表", "合并现金流量表"]) and not is_note_context:
        score += 80
    elif any(token in normalized for token in ["资产负债表", "利润表", "现金流量表"]) and not is_note_context:
        score += 45
    if any(token in normalized for token in ["母公司资产负债表", "母公司利润表", "母公司现金流量表"]):
        score -= 70
    if any(token in normalized for token in ["主要会计数据", "主要财务指标"]):
        score -= 15
    if is_note_context:
        score -= 60
    return score


def _statement_title_position(text: str) -> tuple[str, int] | None:
    normalized = _normalize(text)
    found: list[tuple[int, str]] = []
    for title in STATEMENT_TITLES:
        pos = normalized.find(title)
        if pos >= 0:
            found.append((pos, title))
    if not found:
        return None
    pos, title = sorted(found, key=lambda item: item[0])[0]
    return title, pos


def _row_score(row: SourceRow, canonical: str) -> int:
    score = 50 + _context_score(row.context)
    normalized_context = _normalize(row.context)
    if "table" in row.source:
        score += 25
    if "text" in row.source:
        score += 5
    if "合并" in _normalize(row.context):
        score += 10
    if canonical in BALANCE_SHEET_ITEMS:
        if "资产负债表" in normalized_context and "项目注释" not in normalized_context:
            score += 50
        if "利润表" in normalized_context or "现金流量表" in normalized_context:
            score -= 35
    elif canonical in INCOME_STATEMENT_ITEMS:
        if "利润表" in normalized_context and "项目注释" not in normalized_context:
            score += 50
        if "资产负债表" in normalized_context or "现金流量表补充资料" in normalized_context:
            score -= 35
    elif canonical in CASHFLOW_ITEMS:
        if "现金流量表" in normalized_context and "补充资料" not in normalized_context:
            score += 50
        if "现金流量表补充资料" in normalized_context or "将净利润调节为经营活动现金流量" in normalized_context:
            score -= 80
    row_text = _normalize(row.row_text)
    if row_text.startswith("其中") or "其中:" in row_text:
        score -= 5
    if canonical in {"total_assets", "total_liabilities", "total_equity"} and "合计" in row_text:
        score += 8
    if canonical == "total_equity":
        if row_text.startswith("所有者权益合计") or row_text.startswith("股东权益合计"):
            score += 15
        if "归属于母公司" in row_text or "少数股东" in row_text:
            score -= 10
    return score


def _looks_like_narrative_or_heading(row: SourceRow) -> bool:
    normalized = _normalize(row.row_text)
    display = _display_text(row.row_text)
    numbers = _numbers_from_cells(row.cells)
    if not normalized:
        return True
    if re.match(r"^\(?\d+\)?[、.．]", display) and len(numbers) <= 1:
        return True
    if ("%" in normalized or "％" in normalized) and not any("," in cell or "，" in cell for cell in row.cells):
        return True
    if len(numbers) <= 1 and any(
        token in normalized
        for token in [
            "主要系",
            "较上期",
            "较上年",
            "同比",
            "占公司",
            "占企业",
            "超过",
            "构成",
            "具体为",
            "保持",
            "实现",
            "情况",
            "分析",
            "目录",
            "索引",
        ]
    ):
        return True
    return False


def _fitz_page_texts(path: Path) -> list[str]:
    try:
        import fitz

        with fitz.open(path) as doc:
            return [page.get_text("text") or "" for page in doc]
    except Exception:
        return []


def _rows_from_pdf(path: Path) -> tuple[list[SourceRow], list[str], dict[str, Any]]:
    import pdfplumber

    rows: list[SourceRow] = []
    warnings: list[str] = []
    seen: set[tuple[int | None, str, tuple[str, ...]]] = set()
    stats: dict[str, Any] = {
        "pages_total": 0,
        "pages_scanned": 0,
        "tables_found": 0,
        "table_rows": 0,
        "text_lines": 0,
        "scan_mode": "全文扫描",
    }

    fitz_texts = _fitz_page_texts(path)
    with pdfplumber.open(path) as pdf:
        stats["pages_total"] = len(pdf.pages)
        max_pages = _max_pdf_pages()
        selected_pages = pdf.pages[:max_pages] if max_pages else pdf.pages
        stats["pages_scanned"] = len(selected_pages)
        if max_pages and len(pdf.pages) > max_pages:
            warnings.append(f"年报共 {len(pdf.pages)} 页，本次按 FIN_AGENT_REPORT_MAX_PAGES 限制扫描前 {max_pages} 页。")

        page_texts: list[str] = []
        for page_index, page in enumerate(selected_pages):
            try:
                page_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            except Exception:
                page_text = ""
            if not page_text and page_index < len(fitz_texts):
                page_text = fitz_texts[page_index]
            page_texts.append(page_text)

        active_statement = ""
        for page_index, page in enumerate(selected_pages):
            page_number = page.page_number
            page_text = page_texts[page_index]
            title_position = _statement_title_position(page_text)
            if title_position and title_position[1] < 200:
                active_statement = title_position[0]
            page_context = "\n".join([active_statement, *page_texts[max(0, page_index - 1) : page_index + 1]])

            for settings_index, settings in enumerate(PDF_TABLE_SETTINGS):
                try:
                    tables = page.extract_tables(table_settings=settings) if settings else page.extract_tables()
                except Exception as exc:
                    warnings.append(f"第 {page_number} 页表格策略 {settings_index + 1} 抽取失败：{exc}")
                    continue
                stats["tables_found"] += len(tables or [])
                for table in tables or []:
                    for raw_row in table:
                        cells = [_normalize(cell) for cell in raw_row if _normalize(cell)]
                        if len(cells) == 1:
                            cells = [_normalize(cell) for cell in _split_text_line(cells[0]) if _normalize(cell)]
                        if not cells:
                            continue
                        key = (page_number, f"table:{settings_index}", tuple(cells))
                        if key in seen:
                            continue
                        seen.add(key)
                        rows.append(SourceRow(cells, page_number, f"pdf_table_strategy_{settings_index + 1}", page_context))
                        stats["table_rows"] += 1

            for line in page_text.splitlines():
                cells = _split_text_line(line)
                if len(cells) < 2:
                    continue
                key = (page_number, "text", tuple(cells))
                if key in seen:
                    continue
                seen.add(key)
                rows.append(SourceRow([_normalize(cell) for cell in cells], page_number, "pdf_text_line", page_context))
                stats["text_lines"] += 1

            if title_position and title_position[1] >= 200:
                active_statement = title_position[0]

    if not rows:
        warnings.append("未从 PDF 中识别到可用表格或文本行；扫描版年报可能需要先 OCR。")
    elif stats["text_lines"] == 0 and stats["table_rows"] == 0:
        warnings.append("PDF 可读文本很少；若是扫描版或图片型年报，建议先 OCR。")
    return rows, warnings, stats


def _rows_from_docx(path: Path) -> tuple[list[SourceRow], list[str], dict[str, Any]]:
    from docx import Document

    rows: list[SourceRow] = []
    warnings: list[str] = []
    doc = Document(path)
    paragraphs = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
    stats = {"pages_total": None, "pages_scanned": None, "tables_found": len(doc.tables), "table_rows": 0, "text_lines": 0}

    for table_index, table in enumerate(doc.tables):
        for row in table.rows:
            cells = [_normalize(cell.text) for cell in row.cells if _normalize(cell.text)]
            if cells:
                rows.append(SourceRow(cells, None, f"docx_table_{table_index + 1}", paragraphs))
                stats["table_rows"] += 1

    for paragraph in doc.paragraphs:
        cells = [_split_text_line(paragraph.text)]
        flat_cells = [cell for group in cells for cell in group]
        if len(flat_cells) >= 2:
            rows.append(SourceRow([_normalize(cell) for cell in flat_cells], None, "docx_text_line", paragraphs))
            stats["text_lines"] += 1

    if not rows:
        warnings.append("未从 Word 中识别到表格或可解析文本行。")
    return rows, warnings, stats


def _add_candidate(
    records: dict[str, dict[int, CandidateValue]],
    canonical: str,
    year: int,
    value: float,
    row: SourceRow,
) -> None:
    score = _row_score(row, canonical)
    current = records.setdefault(canonical, {}).get(year)
    candidate = CandidateValue(value=value, score=score, page=row.page, source=row.source, row_text=row.row_text)
    if current is None or candidate.score > current.score:
        records[canonical][year] = candidate


def _extract_records(rows: list[SourceRow], report_year: int | None = None) -> tuple[pd.DataFrame, list[str], dict[str, Any]]:
    aliases = _alias_map()
    records: dict[str, dict[int, CandidateValue]] = {}
    warnings: list[str] = []
    current_year_columns_by_scope: dict[tuple[int | None, str], dict[int, int]] = {}
    current_year_columns_by_source: dict[str, dict[int, int]] = {}

    for row in rows:
        if not row.cells or not any(row.cells):
            continue
        scope = (row.page, row.source)
        year_columns = _years_in_row(row.cells, row.context, report_year)
        canonical = _match_item(row.row_text, aliases)

        if year_columns and not canonical:
            current_year_columns_by_scope[scope] = year_columns
            current_year_columns_by_source[row.source] = year_columns
            continue

        if not canonical:
            continue
        if row.source.startswith("pdf") and _looks_like_narrative_or_heading(row):
            continue

        if year_columns:
            current_year_columns_by_scope[scope] = year_columns
            current_year_columns_by_source[row.source] = year_columns
        current_year_columns = current_year_columns_by_scope.get(scope) or current_year_columns_by_source.get(row.source, {})

        assigned = False
        for col_idx, year in current_year_columns.items():
            if col_idx >= len(row.cells):
                continue
            value = _number(row.cells[col_idx])
            if value is not None:
                _add_candidate(records, canonical, year, value, row)
                assigned = True

        if not assigned and current_year_columns:
            ordered_years = [year for _, year in sorted(current_year_columns.items(), key=lambda item: item[0])]
            numbers = _numbers_from_cells(row.cells[1:])
            for year, value in zip(ordered_years, numbers):
                _add_candidate(records, canonical, year, value, row)
                assigned = True

        if not assigned:
            context_years = _context_years(row.context, report_year)
            numbers = _numbers_from_cells(row.cells)
            if context_years and numbers:
                _add_candidate(records, canonical, max(context_years), numbers[-1], row)

    if not records:
        raise ValueError("未能从年报中识别出标准财务科目。请确认 PDF/Word 包含可复制的财务报表，扫描版 PDF 需要先 OCR。")

    if report_year:
        for canonical, values in list(records.items()):
            for year in list(values):
                if year > report_year or year < report_year - 6:
                    del values[year]
            if not values:
                del records[canonical]
    if not records:
        raise ValueError("识别到的候选数据均不属于报告年度范围，请检查年报是否为可复制财务报表或调整文件。")

    years_before_filter = sorted({year for values in records.values() for year in values})
    core_count_by_year = {
        year: sum(1 for item in REQUIRED_CORE_ITEMS if year in records.get(item, {}))
        for year in years_before_filter
    }
    if len(years_before_filter) > 2:
        minimum_core_count = max(5, int(len(REQUIRED_CORE_ITEMS) * 0.6))
        keep_years = {year for year, count in core_count_by_year.items() if count >= minimum_core_count}
        dropped_years = [year for year in years_before_filter if year not in keep_years]
        if keep_years:
            for values in records.values():
                for year in list(values):
                    if year not in keep_years:
                        del values[year]
            if dropped_years:
                warnings.append(
                    "已剔除核心科目覆盖不足的年度："
                    + "、".join(str(year) for year in dropped_years)
                    + "。这些年份多来自摘要或附注，未进入标准 CSV。"
                )

    years = sorted({year for values in records.values() for year in values})
    output_rows: list[dict[str, Any]] = []
    evidence: dict[str, dict[str, Any]] = {}
    for canonical, values in records.items():
        display_name = DISPLAY_ITEMS.get(canonical, canonical)
        output_row: dict[str, Any] = {"项目": display_name}
        evidence[display_name] = {}
        for year in years:
            candidate = values.get(year)
            output_row[year] = candidate.value if candidate else None
            if candidate:
                evidence[display_name][str(year)] = {
                    "value": candidate.value,
                    "page": candidate.page,
                    "source": candidate.source,
                    "score": candidate.score,
                    "row_text": candidate.row_text[:300],
                }
        output_rows.append(output_row)

    result = pd.DataFrame(output_rows)
    if years:
        result = result[["项目", *years]]

    try:
        cleaned = clean_financial_data(result)
        missing = [item for item in REQUIRED_CORE_ITEMS if item in cleaned.missing_core_items]
        if missing:
            readable = "、".join(DISPLAY_ITEMS.get(item, item) for item in missing)
            warnings.append(f"提取后仍缺少核心科目：{readable}")
    except Exception as exc:
        missing = REQUIRED_CORE_ITEMS
        warnings.append(f"提取结果尚未完全满足标准清洗要求：{exc}")

    core_found = sorted(set(records) & set(REQUIRED_CORE_ITEMS))
    audit = {
        "evidence": evidence,
        "core_items_found": [DISPLAY_ITEMS.get(item, item) for item in core_found],
        "missing_core_items": [DISPLAY_ITEMS.get(item, item) for item in missing],
        "coverage_ratio": round(len(core_found) / len(REQUIRED_CORE_ITEMS), 4),
        "candidate_rows": len(rows),
        "report_year": report_year,
    }
    return result, warnings, audit


def extract_financial_data_from_report(source_path: str | Path, output_dir: str | Path) -> dict[str, Any]:
    path = Path(source_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在：{path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        rows, warnings, stats = _rows_from_pdf(path)
    elif suffix == ".docx":
        rows, warnings, stats = _rows_from_docx(path)
    elif suffix == ".doc":
        raise ValueError("暂不支持旧版二进制 .doc 文件，请先另存为 .docx 后上传。")
    else:
        raise ValueError("仅支持从 PDF、Word(docx) 年报中提取数据。")

    report_year = _infer_report_year(path, rows)
    extracted, parse_warnings, audit = _extract_records(rows, report_year)
    warnings.extend(parse_warnings)
    audit.update(stats)

    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    csv_path = out_dir / "extracted_financial_data.csv"
    xlsx_path = out_dir / "extracted_financial_data.xlsx"
    audit_path = out_dir / "extracted_financial_data_audit.json"
    extracted.to_csv(csv_path, index=False, encoding="utf-8-sig")
    extracted.to_excel(xlsx_path, index=False)
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    years = [int(col) for col in extracted.columns if isinstance(col, int)]

    return {
        "dataframe": extracted,
        "csv_path": csv_path,
        "xlsx_path": xlsx_path,
        "audit_path": audit_path,
        "rows_scanned": len(rows),
        "items_extracted": len(extracted),
        "years": years,
        "warnings": warnings,
        "source_path": path,
        "pages_total": audit.get("pages_total"),
        "pages_scanned": audit.get("pages_scanned"),
        "tables_found": audit.get("tables_found"),
        "table_rows": audit.get("table_rows"),
        "text_lines": audit.get("text_lines"),
        "coverage_ratio": audit.get("coverage_ratio"),
        "missing_core_items": audit.get("missing_core_items", []),
        "evidence": audit.get("evidence", {}),
        "report_year": audit.get("report_year"),
    }

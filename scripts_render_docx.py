from __future__ import annotations

import argparse
import html
import re
from pathlib import Path

import fitz
from docx import Document


PAGE_RECT = fitz.Rect(0, 0, 612, 792)
CONTENT_RECT = fitz.Rect(72, 72, 540, 720)


def _font_archive() -> fitz.Archive | None:
    for path in [
        Path("/System/Library/Fonts"),
        Path("/System/Library/Fonts/Supplemental"),
        Path("/Library/Fonts"),
    ]:
        if path.exists():
            return fitz.Archive(str(path))
    return None


def _style_name(paragraph) -> str:
    return (paragraph.style.name if paragraph.style is not None else "").strip()


def _run_html(paragraph) -> str:
    chunks: list[str] = []
    for run in paragraph.runs:
        text = html.escape(run.text or "")
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        chunks.append(text)
    if chunks:
        return "".join(chunks)
    return html.escape(paragraph.text or "")


def _paragraph_html(paragraph) -> str:
    text = _run_html(paragraph).strip()
    if not text:
        return ""
    style = _style_name(paragraph).lower()
    alignment = getattr(paragraph.paragraph_format, "alignment", None)
    css_class = "center" if str(alignment).endswith("CENTER (1)") else ""

    if "heading 1" in style or style in {"标题 1", "标题1"}:
        return f"<h1>{text}</h1>"
    if "heading 2" in style or style in {"标题 2", "标题2"}:
        return f"<h2>{text}</h2>"
    if "heading 3" in style or style in {"标题 3", "标题3"}:
        return f"<h3>{text}</h3>"
    if text.startswith(("• ", "- ")):
        item = re.sub(r"^(•|-)\s*", "", text)
        return f"<ul><li>{item}</li></ul>"
    return f'<p class="{css_class}">{text}</p>' if css_class else f"<p>{text}</p>"


def _table_html(table) -> str:
    rows: list[str] = []
    for idx, row in enumerate(table.rows):
        tag = "th" if idx == 0 else "td"
        cells = "".join(f"<{tag}>{html.escape(cell.text.strip())}</{tag}>" for cell in row.cells)
        rows.append(f"<tr>{cells}</tr>")
    return "<table>" + "".join(rows) + "</table>"


def docx_to_html(docx_path: Path) -> str:
    doc = Document(docx_path)
    blocks: list[str] = []
    table_iter = iter(doc.tables)
    for child in doc.element.body:
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = None
            for p in doc.paragraphs:
                if p._p is child:
                    paragraph = p
                    break
            if paragraph is not None:
                rendered = _paragraph_html(paragraph)
                if rendered:
                    blocks.append(rendered)
        elif tag == "tbl":
            try:
                blocks.append(_table_html(next(table_iter)))
            except StopIteration:
                continue

    css = """
body {
  font-family: "STHeiti", "STHeiti Medium", "PingFang SC", "Heiti SC", sans-serif;
  font-size: 11pt;
  line-height: 1.35;
  color: #222;
}
h1 {
  color: #1F4D3A;
  font-size: 17pt;
  margin: 15pt 0 8pt 0;
}
h2 {
  color: #2E74B5;
  font-size: 14pt;
  margin: 12pt 0 6pt 0;
}
h3 {
  color: #1F4D78;
  font-size: 12pt;
  margin: 9pt 0 4pt 0;
}
p {
  margin: 0 0 7pt 0;
}
.center {
  text-align: center;
}
ul {
  margin: 0 0 6pt 18pt;
}
li {
  margin-bottom: 3pt;
}
table {
  border-collapse: collapse;
  width: 100%;
  margin: 8pt 0 10pt 0;
}
th, td {
  border: 0.7pt solid #DADCE0;
  padding: 5pt 6pt;
  vertical-align: top;
}
th {
  background: #F2F4F7;
  font-weight: bold;
}
"""
    body = "\n".join(blocks)
    return f"<html><head><meta charset='utf-8'><style>{css}</style></head><body>{body}</body></html>"


def html_to_pdf(html_text: str, pdf_path: Path) -> Path:
    story = fitz.Story(html_text, user_css="", archive=_font_archive())
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    writer = fitz.DocumentWriter(str(pdf_path))

    def rectfn(rect_num: int, filled: bool):
        return PAGE_RECT, CONTENT_RECT, fitz.Matrix(1, 1)

    story.write(writer, rectfn)
    writer.close()
    return pdf_path


def pdf_to_pngs(pdf_path: Path, output_dir: Path, dpi: int) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(pdf_path)
    matrix = fitz.Matrix(dpi / 72, dpi / 72)
    pngs: list[Path] = []
    for index, page in enumerate(doc, start=1):
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        path = output_dir / f"page-{index:03d}.png"
        pix.save(path)
        pngs.append(path)
    return pngs


def render_docx(docx_path: str | Path, output_dir: str | Path | None = None, dpi: int = 144) -> dict[str, object]:
    src = Path(docx_path).expanduser().resolve()
    if not src.exists():
        raise FileNotFoundError(src)
    out_dir = Path(output_dir) if output_dir else Path("outputs") / "rendered_docs" / src.stem
    out_dir.mkdir(parents=True, exist_ok=True)
    html_text = docx_to_html(src)
    html_path = out_dir / f"{src.stem}.html"
    pdf_path = out_dir / f"{src.stem}.pdf"
    html_path.write_text(html_text, encoding="utf-8")
    html_to_pdf(html_text, pdf_path)
    png_paths = pdf_to_pngs(pdf_path, out_dir, dpi)
    return {
        "source": src,
        "output_dir": out_dir,
        "html_path": html_path,
        "pdf_path": pdf_path,
        "png_paths": png_paths,
        "page_count": len(png_paths),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Render DOCX to HTML, PDF and PNG previews without LibreOffice.")
    parser.add_argument("docx", help="DOCX file path")
    parser.add_argument("--output-dir", default=None, help="Output directory")
    parser.add_argument("--dpi", type=int, default=144, help="PNG render DPI")
    args = parser.parse_args()
    result = render_docx(args.docx, args.output_dir, args.dpi)
    print(f"source={result['source']}")
    print(f"output_dir={result['output_dir']}")
    print(f"pdf={result['pdf_path']}")
    print(f"pages={result['page_count']}")


if __name__ == "__main__":
    main()
